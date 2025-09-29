"""
Pathfinder - Modern Streamlit Application
A local AI-powered business intelligence tool with advanced data analysis capabilities
"""

import streamlit as st
import pandas as pd
import numpy as np
import duckdb
import json
import os
from pathlib import Path
import plotly.express as px
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import io
import re
from typing import Dict, List, Optional, Tuple

# PDF and Word document generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Optional import for ER diagrams
try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False

from modules.data_processor import DataProcessor
from modules.visualizer import Visualizer
from modules.llm_agent import LLMAgent
from modules.schema_manager import SchemaManager

# Import Ollama configuration
try:
    from config.ollama_config import ollama_config
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False
    ollama_config = None

# Page configuration
st.set_page_config(
        page_title="Pathfinder",
    page_icon="📊",
    layout="wide",
        initial_sidebar_state="collapsed"
)

    # Fix container width and logo
st.markdown("""
<style>
    .stApp > div {
        padding-top: 0rem;
    }
    
    /* Force 90% width for main container */
    .main .block-container {
        max-width: 90% !important;
        margin: 0 auto !important;
        padding-left: 1rem !important;
        padding-right: 1rem !important;
    }
    
    /* Override all Streamlit containers */
    .stApp > div[data-testid="stAppViewContainer"] {
        max-width: 90% !important;
        margin: 0 auto !important;
    }
    
    .stApp > div[data-testid="stAppViewContainer"] > div[data-testid="stAppViewBlockContainer"] {
        max-width: 90% !important;
        margin: 0 auto !important;
    }
    
    /* Fix logo visibility - force the emoji to show */
    .stApp > div[data-testid="stAppViewContainer"] > div[data-testid="stHeader"] {
        padding-left: 1rem !important;
        padding-right: 1rem !important;
    }
    
    /* Ensure page icon shows properly */
    .stApp > div[data-testid="stAppViewContainer"] > div[data-testid="stHeader"] img {
        display: block !important;
        visibility: visible !important;
    }
    </style>
    """, unsafe_allow_html=True)

# Modern CSS for minimalist design
st.markdown("""
<style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Global Styles */
    .main {
        font-family: 'Inter', sans-serif;
        padding: 0;
        margin: 0;
    }
    
    /* 90% width container with proper centering */
    .main .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
        padding-left: 1rem;
        padding-right: 1rem;
        max-width: 90%;
        margin: 0 auto;
    }
    
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1a1a1a;
        text-align: center;
        margin-bottom: 0.5rem;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        display: inline;
    }
    
    /* Ensure logo emoji is always visible */
    .logo-emoji {
        font-size: 3rem;
        margin-right: 10px;
        display: inline-block;
        vertical-align: middle;
    }
    
    .sub-header {
        font-size: 1.1rem;
        color: #6b7280;
        text-align: center;
        margin-bottom: 2rem;
    }
    
    /* Tab Styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 4px;
        background-color: #f8fafc;
        padding: 6px;
        border-radius: 12px;
        margin-bottom: 2rem;
        overflow-x: auto;
        overflow-y: hidden;
        white-space: nowrap;
        display: flex;
        flex-wrap: nowrap;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: transparent;
        border-radius: 8px;
        padding: 8px 16px;
        font-weight: 500;
        color: #6b7280;
        transition: all 0.2s ease;
        white-space: nowrap;
        flex-shrink: 0;
        min-width: fit-content;
        font-size: 14px;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: #ffffff;
        color: #1a1a1a;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    /* Hide scrollbar but keep functionality */
    .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar {
        height: 4px;
    }
    .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 2px;
    }
    .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar-thumb {
        background: #cbd5e0;
        border-radius: 2px;
    }
    .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar-thumb:hover {
        background: #a0aec0;
    }
    
    /* Data editor enhancements */
    .stDataEditor {
        position: relative;
    }
    
    /* Column header hover effects */
    .stDataEditor [data-testid="stDataEditor"] thead th {
        position: relative;
        cursor: pointer;
    }
    
    .stDataEditor [data-testid="stDataEditor"] thead th:hover {
        background-color: #f0f9ff !important;
    }
    
    /* Add column button on hover for last column */
    .stDataEditor [data-testid="stDataEditor"] thead th:last-child:hover::after {
        content: "➕";
        position: absolute;
        right: 5px;
        top: 50%;
        transform: translateY(-50%);
        font-size: 16px;
        color: #3b82f6;
        pointer-events: none;
    }
    
    /* Right-click context menu styling */
    .column-context-menu {
        position: absolute;
        background: white;
        border: 1px solid #e5e7eb;
        border-radius: 6px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        padding: 4px 0;
        z-index: 1000;
        min-width: 120px;
    }
    
    .column-context-menu-item {
        padding: 8px 16px;
        cursor: pointer;
        color: #374151;
        font-size: 14px;
    }
    
    .column-context-menu-item:hover {
        background-color: #f3f4f6;
    }
    
    .column-context-menu-item.danger {
        color: #dc2626;
    }
    
    .column-context-menu-item.danger:hover {
        background-color: #fef2f2;
    }
    
    /* Card Styling */
    .metric-card {
        background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
        padding: 1.5rem;
        border-radius: 16px;
        border: 1px solid #e2e8f0;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        text-align: center;
        transition: transform 0.2s ease;
    }
    
    .metric-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 15px -3px rgba(0, 0, 0, 0.1);
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: 700;
        color: #1a1a1a;
        margin-bottom: 0.5rem;
    }
    
    .metric-label {
        font-size: 0.875rem;
        color: #6b7280;
        font-weight: 500;
    }
    
    /* Status Indicators */
    .status-success {
        background: linear-gradient(135deg, #10b981 0%, #059669 100%);
        color: white;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        font-weight: 500;
        margin: 1rem 0;
    }
    
    .status-warning {
        background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%);
        color: white;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        font-weight: 500;
        margin: 1rem 0;
    }
    
    .status-error {
        background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%);
        color: white;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        font-weight: 500;
        margin: 1rem 0;
    }
    
    /* Button Styling */
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.75rem 1.5rem;
        font-weight: 500;
        font-size: 0.875rem;
        transition: all 0.2s ease;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.15);
    }
    
    /* File Upload Styling */
    .uploadedFile {
        background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
        border: 2px dashed #cbd5e1;
        border-radius: 12px;
        padding: 2rem;
        text-align: center;
        margin: 1rem 0;
    }
    
    /* Sidebar Styling */
    .css-1d391kg {
        background: linear-gradient(180deg, #f8fafc 0%, #ffffff 100%);
    }
    
    /* Data Table Styling */
    .dataframe {
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    /* Chart Container */
    .chart-container {
        background: white;
        border-radius: 12px;
        padding: 1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        margin: 1rem 0;
    }
    
    /* ER Diagram Container */
    .er-diagram {
        background: white;
        border-radius: 12px;
        padding: 2rem;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        margin: 1rem 0;
        text-align: center;
    }
    
    /* Loading Spinner */
    .stSpinner > div {
        border-top-color: #667eea;
    }
</style>
""", unsafe_allow_html=True)

def get_available_models():
    """Get list of available Ollama models"""
    if not OLLAMA_AVAILABLE:
        return []
    
    try:
        return ollama_config.get_available_models()
    except Exception as e:
        st.error(f"Error getting available models: {str(e)}")
        return []

def get_recommended_models():
    """Get list of recommended models for download"""
    return [
        "llama3:8b-instruct",
        "llama3.2:3b-instruct", 
        "llama3.2:1b-instruct",
        "codellama:7b-instruct",
        "codellama:13b-instruct",
        "mistral:7b-instruct",
        "phi3:3.8b-instruct",
        "deepseek-r1:8b",
        "qwen2.5:7b-instruct",
        "gemma2:9b-instruct"
    ]

def get_missing_models():
    """Get list of recommended models that are not currently available"""
    available = get_available_models()
    recommended = get_recommended_models()
    return [model for model in recommended if model not in available]

def generate_download_commands(models):
    """Generate download commands for the specified models"""
    if not models:
        return []
    
    commands = []
    for model in models:
        commands.append(f"ollama pull {model}")
    
    return commands

def test_model_detection():
    """Test model detection and return status"""
    try:
        available_models = get_available_models()
        return {
            "status": "success",
            "models_found": len(available_models),
            "models": available_models,
            "ollama_available": OLLAMA_AVAILABLE
        }
    except Exception as e:
        return {
            "status": "error",
            "error": str(e),
            "ollama_available": OLLAMA_AVAILABLE
        }

def download_model(model_name: str):
    """Download an Ollama model"""
    with st.spinner(f"Downloading {model_name}..."):
        try:
            import subprocess
            result = subprocess.run(
                ["ollama", "pull", model_name],
                capture_output=True,
                text=True,
                timeout=300
            )
            if result.returncode == 0:
                st.success(f"✅ {model_name} downloaded successfully!")
                st.rerun()
            else:
                st.error(f"❌ Failed to download {model_name}: {result.stderr}")
        except subprocess.TimeoutExpired:
            st.error("❌ Download timeout. Large models may take longer.")
        except Exception as e:
            st.error(f"❌ Error downloading model: {str(e)}")

def safe_dataframe_display(df, width='stretch'):
    """Display DataFrame with proper data type handling to avoid Arrow serialization issues"""
    try:
        # Create a copy and convert problematic data types
        display_df = df.copy()
        for col in display_df.columns:
            if display_df[col].dtype == 'object':
                # Convert object columns to string, handling mixed types
                display_df[col] = display_df[col].astype(str)
            elif 'datetime' in str(display_df[col].dtype):
                # Convert datetime columns to string
                display_df[col] = display_df[col].dt.strftime('%Y-%m-%d %H:%M:%S')
            elif 'Timestamp' in str(display_df[col].dtype):
                # Handle pandas Timestamp objects
                display_df[col] = display_df[col].astype(str)
        
        return st.dataframe(display_df, width=width)
    except Exception as e:
        st.error(f"Error displaying data: {str(e)}")
        # Fallback: display as text
        st.text(str(df.head()))

def sanitize_name(name: str) -> str:
    """Sanitize table/column name for SQL compatibility by replacing spaces with underscores"""
    if not name:
        return name
    
    # Replace spaces with underscores
    sanitized = name.strip().replace(' ', '_')
    
    # Remove any other problematic characters
    sanitized = re.sub(r'[^a-zA-Z0-9_]', '_', sanitized)
    
    # Ensure it starts with letter or underscore
    if sanitized and not re.match(r'^[a-zA-Z_]', sanitized):
        sanitized = f"_{sanitized}"
    
    # Remove consecutive underscores
    sanitized = re.sub(r'_+', '_', sanitized)
    
    # Remove leading/trailing underscores
    sanitized = sanitized.strip('_')
    
    return sanitized

def validate_table_name(name: str) -> Tuple[bool, str]:
    """Validate table name for SQL compatibility"""
    if not name or not name.strip():
        return False, "Table name cannot be empty"
    
    name = name.strip()
    
    # Check for spaces and suggest sanitization
    if ' ' in name:
        sanitized = sanitize_name(name)
        return False, f"Table name cannot contain spaces. Use '{sanitized}' instead"
    
    # Check for SQL reserved words
    sql_reserved = {
        'select', 'from', 'where', 'insert', 'update', 'delete', 'create', 'drop', 'alter',
        'table', 'database', 'index', 'view', 'procedure', 'function', 'trigger', 'constraint',
        'primary', 'foreign', 'key', 'unique', 'check', 'default', 'null', 'not', 'and', 'or',
        'as', 'in', 'like', 'between', 'is', 'exists', 'all', 'any', 'some', 'union', 'intersect',
        'except', 'order', 'group', 'having', 'limit', 'offset', 'distinct', 'top', 'case',
        'when', 'then', 'else', 'end', 'if', 'while', 'for', 'do', 'begin', 'end', 'return',
        'break', 'continue', 'goto', 'declare', 'set', 'exec', 'execute', 'sp_', 'xp_'
    }
    
    if name.lower() in sql_reserved:
        return False, f"'{name}' is a SQL reserved word"
    
    # Check for valid characters (alphanumeric and underscore only)
    if not re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*$', name):
        return False, "Table name must start with letter or underscore and contain only letters, numbers, and underscores"
    
    # Check length
    if len(name) > 128:
        return False, "Table name must be 128 characters or less"
    
    return True, "Valid table name"

def validate_column_name(name: str) -> Tuple[bool, str]:
    """Validate column name for SQL compatibility"""
    if not name or not name.strip():
        return False, "Column name cannot be empty"
    
    name = name.strip()
    
    # Check for spaces and suggest sanitization
    if ' ' in name:
        sanitized = sanitize_name(name)
        return False, f"Column name cannot contain spaces. Use '{sanitized}' instead"
    
    # Check for SQL reserved words
    sql_reserved = {
        'select', 'from', 'where', 'insert', 'update', 'delete', 'create', 'drop', 'alter',
        'table', 'database', 'index', 'view', 'procedure', 'function', 'trigger', 'constraint',
        'primary', 'foreign', 'key', 'unique', 'check', 'default', 'null', 'not', 'and', 'or',
        'as', 'in', 'like', 'between', 'is', 'exists', 'all', 'any', 'some', 'union', 'intersect',
        'except', 'order', 'group', 'having', 'limit', 'offset', 'distinct', 'top', 'case',
        'when', 'then', 'else', 'end', 'if', 'while', 'for', 'do', 'begin', 'end', 'return',
        'break', 'continue', 'goto', 'declare', 'set', 'exec', 'execute', 'sp_', 'xp_'
    }
    
    if name.lower() in sql_reserved:
        return False, f"'{name}' is a SQL reserved word"
    
    # Check for valid characters (alphanumeric and underscore only)
    if not re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*$', name):
        return False, "Column name must start with letter or underscore and contain only letters, numbers, and underscores"
    
    # Check length
    if len(name) > 128:
        return False, "Column name must be 128 characters or less"
    
    return True, "Valid column name"

def check_duplicate_table_name(name: str, existing_tables: Dict) -> Tuple[bool, str]:
    """Check if table name already exists"""
    if name in existing_tables:
        return False, f"Table '{name}' already exists"
    return True, "Table name is available"

def validate_data_entry(df, edited_data):
    """Validate data entry in Excel-like editor"""
    validation_errors = []
    
    for col in df.columns:
        col_type = df[col].dtype
        
        for idx, value in enumerate(edited_data[col]):
            if pd.isna(value) or value == '' or value is None:
                continue  # Skip empty values
                
            # Validate based on column type
            if 'int' in str(col_type) or 'float' in str(col_type):
                try:
                    float(value)
                except (ValueError, TypeError):
                    validation_errors.append(f"Row {idx+1}, Column '{col}': '{value}' is not a valid number")
            
            elif 'datetime' in str(col_type):
                try:
                    pd.to_datetime(value)
                except (ValueError, TypeError):
                    validation_errors.append(f"Row {idx+1}, Column '{col}': '{value}' is not a valid date")
            
            elif col_type == 'object':
                # For text columns, check for reasonable length
                if len(str(value)) > 1000:
                    validation_errors.append(f"Row {idx+1}, Column '{col}': Text too long (max 1000 characters)")
    
    return validation_errors

def create_table_on_enter():
    """Handle Enter key press to create table when all columns are filled"""
    # This function will be called when the last column name is entered
    st.session_state.auto_create_table = True

def parse_markdown_to_reportlab(text: str) -> str:
    """
    Parse markdown text and convert to ReportLab XML formatting
    
    Args:
        text: Markdown text string
        
    Returns:
        ReportLab XML formatted string
    """
    if not text:
        return ""
    
    # First, convert markdown formatting before escaping
    # Convert **bold** to <b>bold</b> (ReportLab supports these tags)
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    
    # Convert *italic* to <i>italic</i> (ReportLab supports these tags)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    
    # Convert `code` to <font name="Courier">code</font> (ReportLab supports font tags)
    text = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', text)
    
    # Convert bullet points to proper formatting
    text = re.sub(r'^[\s]*[•\-]\s*', '• ', text, flags=re.MULTILINE)
    
    # Convert numbered lists
    text = re.sub(r'^[\s]*(\d+)\.\s*', r'\1. ', text, flags=re.MULTILINE)
    
    # Now escape special XML characters, but be more careful with dollar signs
    # Only escape standalone dollar signs, not those in mathematical expressions
    # Escape ampersands first to avoid double-escaping
    text = text.replace('&', '&amp;')
    
    # Escape dollar signs more carefully - only if they're not part of LaTeX math
    # This regex looks for dollar signs that are not part of $$...$$ or $...$ patterns
    text = re.sub(r'(?<!\$)\$(?!\$)', '&#36;', text)
    
    # Escape other special XML characters, but preserve our formatting tags
    # We need to be careful not to escape the <b>, </b>, <i>, </i>, <font>, </font> tags
    # First, temporarily replace our formatting tags
    text = text.replace('<b>', '___BOLD_START___')
    text = text.replace('</b>', '___BOLD_END___')
    text = text.replace('<i>', '___ITALIC_START___')
    text = text.replace('</i>', '___ITALIC_END___')
    text = text.replace('<font name="Courier">', '___FONT_START___')
    text = text.replace('</font>', '___FONT_END___')
    
    # Now escape other < and > characters
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    
    # Restore our formatting tags
    text = text.replace('___BOLD_START___', '<b>')
    text = text.replace('___BOLD_END___', '</b>')
    text = text.replace('___ITALIC_START___', '<i>')
    text = text.replace('___ITALIC_END___', '</i>')
    text = text.replace('___FONT_START___', '<font name="Courier">')
    text = text.replace('___FONT_END___', '</font>')
    
    # Convert line breaks to proper spacing
    text = text.replace('\n', '<br/>')
    
    return text

def escape_markdown_for_streamlit(text: str) -> str:
    """
    Escape markdown text for proper display in Streamlit
    
    Args:
        text: Markdown text string
        
    Returns:
        Streamlit-safe markdown string
    """
    if not text:
        return ""
    
    # Escape dollar signs to prevent Streamlit from interpreting them as LaTeX math
    # Use backslash to escape dollar signs in markdown
    text = text.replace('$', '\\$')
    
    return text

def generate_pdf_report(insight_data: Dict, filename: str) -> bytes:
    """Generate PDF report using ReportLab"""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("ReportLab not available. Install with: pip install reportlab")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=80, bottomMargin=50)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Enhanced title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=28,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#2c3e50'),
        fontName='Helvetica-Bold'
    )
    
    # Enhanced heading styles
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=15,
        spaceBefore=20,
        textColor=colors.HexColor('#34495e'),
        fontName='Helvetica-Bold',
        borderWidth=0,
        borderColor=colors.HexColor('#3498db'),
        borderPadding=5,
        backColor=colors.HexColor('#ecf0f1')
    )
    
    # Subheading style
    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.HexColor('#2c3e50'),
        fontName='Helvetica-Bold'
    )
    
    # Enhanced normal style
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=8,
        textColor=colors.HexColor('#2c3e50'),
        fontName='Helvetica',
        leading=16
    )
    
    # Enhanced code style
    code_style = ParagraphStyle(
        'CustomCode',
        parent=styles['Code'],
        fontSize=11,
        leftIndent=15,
        rightIndent=15,
        spaceAfter=15,
        spaceBefore=10,
        backgroundColor=colors.HexColor('#f8f9fa'),
        textColor=colors.HexColor('#2c3e50'),
        fontName='Courier-Bold',
        borderWidth=2,
        borderColor=colors.HexColor('#3498db'),
        borderPadding=12
    )
    
    # Metadata style
    metadata_style = ParagraphStyle(
        'CustomMetadata',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=5,
        textColor=colors.HexColor('#7f8c8d'),
        fontName='Helvetica',
        alignment=TA_CENTER
    )
    
    # Build content
    story = []
    
    # Header with logo and title
    story.append(Paragraph("🤖 AI Analysis Report", title_style))
    story.append(Spacer(1, 10))
    
    # Decorative line
    story.append(Paragraph("─" * 50, metadata_style))
    story.append(Spacer(1, 15))
    
    # Metadata in a nice box
    metadata_table = Table([
        ["Generated:", insight_data['timestamp']],
        ["Data Shape:", f"{insight_data['data_shape'][0]} rows × {insight_data['data_shape'][1]} columns"],
        ["Report Type:", "AI-Powered Business Intelligence Analysis"]
    ], colWidths=[120, 300])
    
    metadata_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7'))
    ]))
    
    story.append(metadata_table)
    story.append(Spacer(1, 25))
    
    # Question section with icon
    story.append(Paragraph("📝 Analysis Question", heading_style))
    story.append(Spacer(1, 10))
    question_text = parse_markdown_to_reportlab(insight_data["question"])
    story.append(Paragraph(f'<i>"{question_text}"</i>', normal_style))
    story.append(Spacer(1, 20))
    
    # SQL Query section
    story.append(Paragraph("🔍 SQL Query Executed", heading_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(insight_data['sql_query'], code_style))
    story.append(Spacer(1, 25))
    
    # AI Analysis section
    story.append(Paragraph("🧠 AI Analysis & Insights", heading_style))
    story.append(Spacer(1, 15))
    
    # Split analysis into paragraphs and format better
    analysis_paragraphs = insight_data['analysis'].split('\n\n')
    for i, para in enumerate(analysis_paragraphs):
        if para.strip():
            # Check if it's a heading (starts with **)
            if para.strip().startswith('**') and para.strip().endswith('**'):
                # It's a heading - clean it and apply heading style
                clean_heading = para.strip().replace('**', '').replace('*', '')
                story.append(Paragraph(clean_heading, subheading_style))
            elif para.strip().startswith('•') or para.strip().startswith('-'):
                # It's a bullet point - parse markdown formatting
                bullet_text = para.strip().lstrip('•- ')
                formatted_text = parse_markdown_to_reportlab(bullet_text)
                story.append(Paragraph(f"• {formatted_text}", normal_style))
            elif para.strip().startswith(tuple('123456789')):
                # It's a numbered list - parse markdown formatting
                formatted_text = parse_markdown_to_reportlab(para.strip())
                story.append(Paragraph(formatted_text, normal_style))
            else:
                # Regular paragraph - parse markdown formatting
                formatted_text = parse_markdown_to_reportlab(para.strip())
                story.append(Paragraph(formatted_text, normal_style))
            story.append(Spacer(1, 8))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_bulk_pdf_report(insights_history: List[Dict]) -> bytes:
    """Generate bulk PDF report with all insights"""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("ReportLab not available. Install with: pip install reportlab")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=80, bottomMargin=50)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Enhanced title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=28,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#2c3e50'),
        fontName='Helvetica-Bold'
    )
    
    # Enhanced heading styles
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=15,
        spaceBefore=20,
        textColor=colors.HexColor('#34495e'),
        fontName='Helvetica-Bold',
        borderWidth=0,
        borderColor=colors.HexColor('#3498db'),
        borderPadding=5,
        backColor=colors.HexColor('#ecf0f1')
    )
    
    # Subheading style
    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.HexColor('#2c3e50'),
        fontName='Helvetica-Bold'
    )
    
    # Enhanced normal style
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=8,
        textColor=colors.HexColor('#2c3e50'),
        fontName='Helvetica',
        leading=16
    )
    
    # Enhanced code style
    code_style = ParagraphStyle(
        'CustomCode',
        parent=styles['Code'],
        fontSize=11,
        leftIndent=15,
        rightIndent=15,
        spaceAfter=15,
        spaceBefore=10,
        backgroundColor=colors.HexColor('#f8f9fa'),
        textColor=colors.HexColor('#2c3e50'),
        fontName='Courier-Bold',
        borderWidth=2,
        borderColor=colors.HexColor('#3498db'),
        borderPadding=12
    )
    
    # Metadata style
    metadata_style = ParagraphStyle(
        'CustomMetadata',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=5,
        textColor=colors.HexColor('#7f8c8d'),
        fontName='Helvetica',
        alignment=TA_CENTER
    )
    
    # Build content
    story = []
    
    # Header with logo and title
    story.append(Paragraph("🤖 AI Insights Report - All Analysis", title_style))
    story.append(Spacer(1, 10))
    
    # Decorative line
    story.append(Paragraph("─" * 50, metadata_style))
    story.append(Spacer(1, 15))
    
    # Summary in a nice table
    summary_table = Table([
        ["Total Insights:", str(len(insights_history))],
        ["Generated:", datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ["Report Type:", "Comprehensive AI Business Intelligence Analysis"]
    ], colWidths=[120, 300])
    
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e74c3c')),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7'))
    ]))
    
    story.append(summary_table)
    story.append(Spacer(1, 25))
    
    # Add each insight
    for i, insight in enumerate(insights_history):
        # Insight header with better styling
        story.append(Paragraph(f"📊 Analysis #{i+1}", heading_style))
        story.append(Spacer(1, 10))
        
        # Metadata in a table
        metadata_table = Table([
            ["Generated:", insight['timestamp']],
            ["Data Shape:", f"{insight['data_shape'][0]} rows × {insight['data_shape'][1]} columns"]
        ], colWidths=[100, 320])
        
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#95a5a6')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7'))
        ]))
        
        story.append(metadata_table)
        story.append(Spacer(1, 15))
        
        # Question section
        story.append(Paragraph("📝 Analysis Question", subheading_style))
        story.append(Spacer(1, 8))
        story.append(Paragraph(f'<i>"{insight["question"]}"</i>', normal_style))
        story.append(Spacer(1, 15))
        
        # SQL Query section
        story.append(Paragraph("🔍 SQL Query Executed", subheading_style))
        story.append(Spacer(1, 8))
        story.append(Paragraph(insight['sql_query'], code_style))
        story.append(Spacer(1, 15))
        
        # AI Analysis section
        story.append(Paragraph("🧠 AI Analysis & Insights", subheading_style))
        story.append(Spacer(1, 10))
        
        # Split analysis into paragraphs and format better
        analysis_paragraphs = insight['analysis'].split('\n\n')
        for para in analysis_paragraphs:
            if para.strip():
                # Check if it's a heading (starts with **)
                if para.strip().startswith('**') and para.strip().endswith('**'):
                    # It's a heading - clean it and apply heading style
                    clean_heading = para.strip().replace('**', '').replace('*', '')
                    story.append(Paragraph(clean_heading, subheading_style))
                elif para.strip().startswith('•') or para.strip().startswith('-'):
                    # It's a bullet point - parse markdown formatting
                    bullet_text = para.strip().lstrip('•- ')
                    formatted_text = parse_markdown_to_reportlab(bullet_text)
                    story.append(Paragraph(f"• {formatted_text}", normal_style))
                elif para.strip().startswith(tuple('123456789')):
                    # It's a numbered list - parse markdown formatting
                    formatted_text = parse_markdown_to_reportlab(para.strip())
                    story.append(Paragraph(formatted_text, normal_style))
                else:
                    # Regular paragraph - parse markdown formatting
                    formatted_text = parse_markdown_to_reportlab(para.strip())
                    story.append(Paragraph(formatted_text, normal_style))
                story.append(Spacer(1, 6))
        
        # Add page break between insights (except for the last one)
        if i < len(insights_history) - 1:
            story.append(PageBreak())
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_word_report(insight_data: Dict) -> bytes:
    """Generate Word document report"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available. Install with: pip install python-docx")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('🤖 AI Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Generated: {insight_data['timestamp']}")
    doc.add_paragraph(f"Data Shape: {insight_data['data_shape'][0]} rows × {insight_data['data_shape'][1]} columns")
    doc.add_paragraph("")
    
    # Question
    doc.add_heading('📝 Question', level=1)
    doc.add_paragraph(insight_data['question'])
    
    # SQL Query
    doc.add_heading('🔍 SQL Query', level=1)
    sql_para = doc.add_paragraph(insight_data['sql_query'])
    sql_para.style = 'Code'
    
    # AI Analysis
    doc.add_heading('🧠 AI Analysis', level=1)
    
    # Split analysis into paragraphs
    analysis_paragraphs = insight_data['analysis'].split('\n\n')
    for para in analysis_paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def handle_tab_navigation():
    """Handle tab navigation requests"""
    if st.session_state.navigate_to_tab:
        tab_name = st.session_state.navigate_to_tab
        st.session_state.navigate_to_tab = None  # Reset after use
        
        # Show prominent notification for the target tab with better styling
        if tab_name == "data_load":
            st.markdown("""
            <div style="background-color: #e3f2fd; border: 2px solid #2196f3; border-radius: 10px; padding: 15px; margin: 10px 0;">
                <h3 style="color: #1976d2; margin: 0;">📁 Navigate to Data Load!</h3>
                <p style="margin: 5px 0 0 0; color: #424242;">Click on the <strong>'📁 Data Load'</strong> tab above to load your data</p>
            </div>
            """, unsafe_allow_html=True)
        elif tab_name == "explorer":
            st.markdown("""
            <div style="background-color: #e8f5e8; border: 2px solid #4caf50; border-radius: 10px; padding: 15px; margin: 10px 0;">
                <h3 style="color: #388e3c; margin: 0;">🔍 Navigate to Explorer!</h3>
                <p style="margin: 5px 0 0 0; color: #424242;">Click on the <strong>'🔍 Explorer'</strong> tab above to analyze your data</p>
            </div>
            """, unsafe_allow_html=True)
        elif tab_name == "relationships":
            st.markdown("""
            <div style="background-color: #fff3e0; border: 2px solid #ff9800; border-radius: 10px; padding: 15px; margin: 10px 0;">
                <h3 style="color: #f57c00; margin: 0;">🔗 Navigate to Relationships!</h3>
                <p style="margin: 5px 0 0 0; color: #424242;">Click on the <strong>'🔗 Relationships'</strong> tab above to define relationships</p>
            </div>
            """, unsafe_allow_html=True)
        elif tab_name == "queries":
            st.markdown("""
            <div style="background-color: #f3e5f5; border: 2px solid #9c27b0; border-radius: 10px; padding: 15px; margin: 10px 0;">
                <h3 style="color: #7b1fa2; margin: 0;">💻 Navigate to SQL Playground!</h3>
                <p style="margin: 5px 0 0 0; color: #424242;">Click on the <strong>'💻 SQL Playground'</strong> tab above to create custom queries</p>
            </div>
            """, unsafe_allow_html=True)
        elif tab_name == "ai":
            st.markdown("""
            <div style="background-color: #e1f5fe; border: 2px solid #00bcd4; border-radius: 10px; padding: 15px; margin: 10px 0;">
                <h3 style="color: #0097a7; margin: 0;">🤖 Navigate to AI Analysis!</h3>
                <p style="margin: 5px 0 0 0; color: #424242;">Click on the <strong>'🤖 AI Analysis'</strong> tab above for intelligent insights</p>
            </div>
            """, unsafe_allow_html=True)
        elif tab_name == "settings":
            st.markdown("""
            <div style="background-color: #fce4ec; border: 2px solid #e91e63; border-radius: 10px; padding: 15px; margin: 10px 0;">
                <h3 style="color: #c2185b; margin: 0;">⚙️ Navigate to Settings!</h3>
                <p style="margin: 5px 0 0 0; color: #424242;">Click on the <strong>'⚙️ Settings'</strong> tab above to configure the system</p>
            </div>
            """, unsafe_allow_html=True)
        elif tab_name == "insights":
            st.markdown("""
            <div style="background-color: #e8f5e8; border: 2px solid #4caf50; border-radius: 10px; padding: 15px; margin: 10px 0;">
                <h3 style="color: #388e3c; margin: 0;">📄 Navigate to Insights!</h3>
                <p style="margin: 5px 0 0 0; color: #424242;">Click on the <strong>'🤖 AI Analysis'</strong> tab above to view generated insights</p>
            </div>
            """, unsafe_allow_html=True)


def initialize_session_state():
    """Initialize all session state variables"""
    if 'data_processor' not in st.session_state:
        st.session_state.data_processor = DataProcessor()
    if 'visualizer' not in st.session_state:
        st.session_state.visualizer = Visualizer()
    if 'llm_agent' not in st.session_state:
        st.session_state.llm_agent = LLMAgent()
    if 'schema_manager' not in st.session_state:
        st.session_state.schema_manager = SchemaManager()
    if 'loaded_tables' not in st.session_state:
        st.session_state.loaded_tables = {}
    if 'selected_columns' not in st.session_state:
        st.session_state.selected_columns = {}
    if 'table_relationships' not in st.session_state:
        st.session_state.table_relationships = []
    if 'relationships' not in st.session_state:
        st.session_state.relationships = []
    if 'ai_analysis_results' not in st.session_state:
        st.session_state.ai_analysis_results = {}
    
    # Initialize AI Analysis related variables
    if 'current_user_query' not in st.session_state:
        st.session_state.current_user_query = ""
    if 'current_sql_query' not in st.session_state:
        st.session_state.current_sql_query = ""
    if 'ai_analysis_sql_query' not in st.session_state:
        st.session_state.ai_analysis_sql_query = ""
    if 'custom_sql_query' not in st.session_state:
        st.session_state.custom_sql_query = ""
    if 'current_query_results' not in st.session_state:
        st.session_state.current_query_results = None
    if 'ai_analysis_query_results' not in st.session_state:
        st.session_state.ai_analysis_query_results = None
    if 'sql_playground_query_results' not in st.session_state:
        st.session_state.sql_playground_query_results = None
    if 'analysis_query' not in st.session_state:
        st.session_state.analysis_query = ""
    if 'generate_question_flag' not in st.session_state:
        st.session_state.generate_question_flag = False
    if 'generated_question' not in st.session_state:
        st.session_state.generated_question = ""
    if 'question_to_display' not in st.session_state:
        st.session_state.question_to_display = ""
    if 'current_visualization' not in st.session_state:
        st.session_state.current_visualization = None
    if 'current_ai_analysis' not in st.session_state:
        st.session_state.current_ai_analysis = ""
    if 'analysis_timestamp' not in st.session_state:
        st.session_state.analysis_timestamp = ""
    
    # Initialize tab navigation
    if 'navigate_to_tab' not in st.session_state:
        st.session_state.navigate_to_tab = None

def load_multi_sheet_excel(file_upload):
    """Load all sheets from Excel file with sanitized table names"""
    try:
        # Read the file into memory first
        file_content = file_upload.read()
        
        # Create ExcelFile object from bytes
        excel_file = pd.ExcelFile(io.BytesIO(file_content))
        sheet_names = excel_file.sheet_names
        
        loaded_tables = {}
        for sheet_name in sheet_names:
            try:
                # Read each sheet from the ExcelFile object
                df = excel_file.parse(sheet_name)
                if not df.empty:
                    # Sanitize sheet name for SQL compatibility
                    sanitized_name = sanitize_name(sheet_name)
                    
                    # Sanitize column names as well
                    df.columns = [sanitize_name(col) for col in df.columns]
                    
                    # Use sanitized name as table name
                    loaded_tables[sanitized_name] = df
                    
                    # Show warning if name was changed
                    if sanitized_name != sheet_name:
                        st.info(f"📝 Sheet '{sheet_name}' renamed to '{sanitized_name}' for SQL compatibility")
            except Exception as e:
                st.warning(f"Could not load sheet '{sheet_name}': {str(e)}")
        
        return loaded_tables
    except Exception as e:
        st.error(f"Error loading Excel file: {str(e)}")
        return {}

def create_er_diagram(tables, relationships):
    """Create a beautiful, large, and readable Entity Relationship diagram using NetworkX and Matplotlib"""
    if not NETWORKX_AVAILABLE:
        st.warning("NetworkX not available. Install with: pip install networkx")
        return None
    
    try:
        G = nx.DiGraph()
        
        # Add nodes for tables with metadata
        for table_name, df in tables.items():
            G.add_node(table_name, node_type='table', columns=df.columns.tolist(), rows=df.shape[0])
        
        # Add edges for relationships
        for rel in relationships:
            if rel['source_table'] in tables and rel['target_table'] in tables:
                G.add_edge(rel['source_table'], rel['target_table'], 
                          label=f"{rel['source_column']} → {rel['target_column']}",
                          rel_type=rel.get('type', 'One-to-Many'))
        
        # Much larger figure size for better readability
        num_tables = len(tables)
        if num_tables <= 2:
            figsize = (20, 12)
        elif num_tables <= 4:
            figsize = (24, 16)
        elif num_tables <= 6:
            figsize = (28, 20)
        else:
            figsize = (32, 24)
        
        # Create the plot with much larger sizing
        fig, ax = plt.subplots(figsize=figsize, dpi=100)
        fig.patch.set_facecolor('white')
        ax.set_facecolor('#f8f9fa')
        
        # Create layout with much better spacing
        pos = nx.spring_layout(G, k=5, iterations=200, seed=42)
        
        # Title with larger font
        ax.set_title("📊 Entity Relationship Diagram", fontsize=28, fontweight='bold', 
                    color='#2c3e50', pad=30)
        
        # Draw nodes with much larger and more beautiful styling
        node_colors = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']
        for i, (node, data) in enumerate(G.nodes(data=True)):
            if data['node_type'] == 'table':
                color = node_colors[i % len(node_colors)]
                
                # Draw the main node circle
                nx.draw_networkx_nodes(G, pos, 
                                      nodelist=[node],
                                      node_color=color, 
                                      node_size=8000,  # Much larger nodes
                                      alpha=0.9,
                                      ax=ax)
                
                # Add table name as main text
                x, y = pos[node]
                ax.text(x, y, f"{node}", 
                       ha='center', va='center', fontsize=16, fontweight='bold', 
                       color='white', bbox=dict(boxstyle="round,pad=0.3", facecolor=color, alpha=0.9))
                
                # Add table info below the name
                ax.text(x, y-0.15, f"({data['rows']} rows, {len(data['columns'])} cols)", 
                       ha='center', va='center', fontsize=12, fontweight='normal', 
                       color='white', bbox=dict(boxstyle="round,pad=0.2", facecolor=color, alpha=0.7))
                
                # Add column names in a smaller box below
                columns_text = ', '.join(data['columns'][:5])  # Show first 5 columns
                if len(data['columns']) > 5:
                    columns_text += f"... (+{len(data['columns'])-5} more)"
                
                ax.text(x, y-0.3, columns_text, 
                       ha='center', va='center', fontsize=10, fontweight='normal', 
                       color='#2c3e50', bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.9, edgecolor=color))
        
        # Draw edges with much better styling
        edge_colors = ['#7f8c8d' for _ in G.edges()]
        nx.draw_networkx_edges(G, pos, 
                              edge_color=edge_colors, 
                              arrows=True, 
                              arrowsize=30,  # Larger arrows
                              alpha=0.8,
                              width=4,  # Thicker lines
                              arrowstyle='->',
                              ax=ax)
        
        # Draw edge labels with much better positioning and styling
        edge_labels = nx.get_edge_attributes(G, 'label')
        for edge, label in edge_labels.items():
            x1, y1 = pos[edge[0]]
            x2, y2 = pos[edge[1]]
            mid_x, mid_y = (x1 + x2) / 2, (y1 + y2) / 2
            
            # Offset the label slightly to avoid overlap with lines
            offset_x = (y2 - y1) * 0.1
            offset_y = -(x2 - x1) * 0.1
            
            ax.text(mid_x + offset_x, mid_y + offset_y, label, ha='center', va='center', 
                   fontsize=12, fontweight='bold', color='#2c3e50',
                   bbox=dict(boxstyle="round,pad=0.3", facecolor='#ecf0f1', alpha=0.9, edgecolor='#bdc3c7', linewidth=2))
        
        # Add a comprehensive legend
        legend_elements = [
            plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#3498db', 
                      markersize=20, label='Tables'),
            plt.Line2D([0], [0], color='#7f8c8d', linewidth=4, label='Relationships'),
            plt.Line2D([0], [0], marker='s', color='w', markerfacecolor='#ecf0f1', 
                      markersize=15, label='Column Mappings')
        ]
        ax.legend(handles=legend_elements, loc='upper right', fontsize=14, framealpha=0.9)
        
        # Add relationship type legend if there are relationships
        if relationships:
            rel_types = set(rel.get('type', 'One-to-Many') for rel in relationships)
            rel_legend_text = f"Relationship Types: {', '.join(rel_types)}"
            ax.text(0.02, 0.98, rel_legend_text, transform=ax.transAxes, 
                   fontsize=12, fontweight='bold', color='#2c3e50',
                   bbox=dict(boxstyle="round,pad=0.3", facecolor='#ecf0f1', alpha=0.9))
        
        # Remove axes and add subtle grid
        ax.axis('off')
        ax.grid(True, alpha=0.1)
        
        # Ensure the diagram fits properly with more padding
        plt.tight_layout(pad=2.0)
        return fig
        
    except Exception as e:
        st.error(f"Error creating ER diagram: {str(e)}")
        return None

def main():
    """Main application function"""
    initialize_session_state()
    
    # Header with visible logo
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown('<div style="text-align: center;"><span class="logo-emoji">📊</span><h1 class="main-header">Pathfinder</h1></div>', unsafe_allow_html=True)
        st.markdown('<p class="sub-header">AI-Powered Data Analysis</p>', unsafe_allow_html=True)
    
    # Handle tab navigation - show notifications after header
    handle_tab_navigation()
    
    # Main tabs - reduced to fit better
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📁 Data Load", 
        "🔍 Explorer", 
        "🔗 Relationships", 
        "💻 SQL Playground",
        "🤖 AI Analysis", 
        "⚙️ Settings"
    ])
    
    with tab1:
        data_load_tab()
    
    with tab2:
        data_explorer_tab()
    
    with tab3:
        relationship_builder_tab()
    
    with tab4:
        custom_queries_tab()
    
    with tab5:
        ai_analysis_tab()
    
    with tab6:
        settings_tab()

def show_table_editor_interface(table_name, df):
    """Show the table editor interface for a specific table"""
    # Initialize edited dataframe if not exists
    if f"edited_df_{table_name}" not in st.session_state:
        st.session_state[f"edited_df_{table_name}"] = df.copy()
        # Add empty rows for new data entry
        empty_rows = pd.DataFrame(index=range(3), columns=df.columns)
        empty_rows = empty_rows.fillna('')  # Fill with empty strings for editing
        st.session_state[f"edited_df_{table_name}"] = pd.concat([st.session_state[f"edited_df_{table_name}"], empty_rows], ignore_index=True)
    
    edited_df = st.session_state[f"edited_df_{table_name}"]
    
    # Show current table info and finish editing option
    col1, col2 = st.columns([3, 1])
    with col1:
        st.info(f"Currently editing: **{table_name}** ({edited_df.shape[0]} rows × {edited_df.shape[1]} columns)")
    with col2:
        if st.button("✅ Finish Editing", key=f"finish_edit_{table_name}"):
            # Save changes before finishing
            non_empty_rows = st.session_state[f"edited_df_{table_name}"][~(st.session_state[f"edited_df_{table_name}"] == '').all(axis=1)]
            if not non_empty_rows.empty:
                # Save to appropriate location
                if table_name in st.session_state.loaded_tables:
                    st.session_state.loaded_tables[table_name] = non_empty_rows
                elif table_name == 'main':
                    st.session_state.data_processor.data = non_empty_rows
                    st.session_state.data_processor._register_data_in_duckdb()
                
                # Clear the edited_df from session state
                if f"edited_df_{table_name}" in st.session_state:
                    del st.session_state[f"edited_df_{table_name}"]
                
                st.success("✅ Changes saved and editing finished!")
            else:
                st.warning("⚠️ No data to save.")
            
            # Exit editing mode but keep expander open
            st.session_state[f"editing_table_{table_name}"] = False
            st.session_state[f"expander_open_{table_name}"] = True
            st.rerun()
    
    # Excel-like interface with column headers and add buttons
    if len(edited_df.columns) > 6:
        st.markdown("**📊 Wide Table View: Click on any cell to edit. Use the + buttons to add columns/rows. Scroll horizontally to see all columns.**")
    else:
        st.markdown("**Click on any cell to edit. Use the + buttons to add columns/rows.**")
    
    # Pagination for large tables
    rows_per_page = 20
    total_rows = len(edited_df)
    
    if total_rows > rows_per_page:
        # Initialize pagination state
        if f"current_page_{table_name}" not in st.session_state:
            st.session_state[f"current_page_{table_name}"] = 1
        
        total_pages = (total_rows + rows_per_page - 1) // rows_per_page  # Ceiling division
        current_page = st.session_state[f"current_page_{table_name}"]
        
        # Pagination controls
        col1, col2, col3, col4, col5 = st.columns([1, 1, 2, 1, 1])
        
        with col1:
            if st.button("⏮️ First", key=f"first_page_{table_name}"):
                st.session_state[f"current_page_{table_name}"] = 1
                st.rerun()
        
        with col2:
            if st.button("⬅️ Prev", key=f"prev_page_{table_name}") and current_page > 1:
                st.session_state[f"current_page_{table_name}"] = current_page - 1
                st.rerun()
        
        with col3:
            st.markdown(f"**Page {current_page} of {total_pages}** ({total_rows} total rows)")
        
        with col4:
            if st.button("➡️ Next", key=f"next_page_{table_name}") and current_page < total_pages:
                st.session_state[f"current_page_{table_name}"] = current_page + 1
                st.rerun()
        
        with col5:
            if st.button("⏭️ Last", key=f"last_page_{table_name}"):
                st.session_state[f"current_page_{table_name}"] = total_pages
                st.rerun()
        
        # Custom page input - better positioned
        st.markdown("---")
        col1, col2, col3, col4 = st.columns([1, 1, 1, 2])
        with col1:
            custom_page = st.number_input(
                "Go to page:",
                min_value=1,
                max_value=total_pages,
                value=current_page,
                key=f"custom_page_{table_name}"
            )
        with col2:
            if st.button("Go", key=f"go_page_{table_name}"):
                st.session_state[f"current_page_{table_name}"] = custom_page
                st.rerun()
        with col3:
            st.markdown("")  # Empty space for alignment
        with col4:
            st.markdown(f"**Showing rows {(current_page-1)*rows_per_page + 1} to {min(current_page*rows_per_page, total_rows)}**")
        
        # Get the current page data
        start_idx = (current_page - 1) * rows_per_page
        end_idx = min(start_idx + rows_per_page, total_rows)
        current_page_data = edited_df.iloc[start_idx:end_idx]
        
        st.markdown("---")
    else:
        # No pagination needed for small tables
        current_page_data = edited_df
        current_page = 1
    
    # Create the table with Excel-like headers (including row number column)
    # Adjust column widths based on number of columns for better readability
    if len(edited_df.columns) > 6:
        # For wide tables, use fixed widths to ensure proper scrolling
        col_widths = [0.08] + [0.92] * len(edited_df.columns) + [0.08]  # Row numbers + data columns + action buttons
    else:
        # For narrow tables, use normal column widths
        col_widths = [0.3] + [1] * len(edited_df.columns) + [0.3]  # Row numbers + data columns + action buttons
    
    # Add horizontal scroll for tables with more than 6 columns
    if len(edited_df.columns) > 6:
        st.markdown("**📊 Wide Table View: Scroll horizontally to see all columns**")
        st.markdown("---")
        
        # Use a more effective CSS approach that works with Streamlit
        st.markdown("""
        <style>
        .wide-table-scroll {
            overflow-x: auto !important;
            border: 2px solid #007bff;
            border-radius: 8px;
            padding: 15px;
            background-color: #f8f9fa;
            margin: 10px 0;
            width: 100% !important;
        }
        .wide-table-scroll .stColumns {
            min-width: max-content !important;
            flex-wrap: nowrap !important;
            width: max-content !important;
        }
        .wide-table-scroll .stTextInput > div > div > input {
            min-width: 120px !important;
            width: 120px !important;
        }
        .wide-table-scroll .stButton > button {
            min-width: 35px !important;
        }
        .wide-table-scroll .stMarkdown {
            white-space: nowrap !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # Start the scrollable container
        st.markdown('<div class="wide-table-scroll">', unsafe_allow_html=True)
    
    # Column headers with add/delete column functionality
    header_cols = st.columns(col_widths)
    
    # Row number header
    with header_cols[0]:
        st.markdown("**#**")
    
    # Data column headers
    for i, col in enumerate(edited_df.columns):
        with header_cols[i + 1]:  # +1 because first column is row numbers
            # Column header with rename, add, and delete options
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col1:
                # Directly editable column name
                new_col_name = st.text_input(
                    "Column name:",
                    value=col,
                    key=f"header_input_{col}_{table_name}",
                    help=f"Edit column name: {col}",
                    label_visibility="collapsed"
                )
                
                # Check if column name was changed
                if new_col_name != col and new_col_name:
                    is_valid, message = validate_column_name(new_col_name)
                    if is_valid:
                        # Rename the column
                        st.session_state[f"edited_df_{table_name}"].rename(columns={col: new_col_name}, inplace=True)
                        df.rename(columns={col: new_col_name}, inplace=True)
                        st.rerun()
                    else:
                        st.error(f"❌ {message}")
            
            with col2:
                if st.button("➕", key=f"add_col_after_{col}_{table_name}", help=f"Add column after {col}"):
                    # Add new column after current column
                    new_col_name = f"Column{len(edited_df.columns) + 1}"
                    new_position = i + 1
                    
                    # Insert new column
                    st.session_state[f"edited_df_{table_name}"].insert(new_position, new_col_name, "")
                    df.insert(new_position, new_col_name, "")
                    # Ensure editing state and expander state are maintained
                    st.session_state[f"editing_table_{table_name}"] = True
                    st.session_state[f"expander_open_{table_name}"] = True
                    st.rerun()
            
            with col3:
                if len(edited_df.columns) > 1:  # Don't allow deleting the last column
                    if st.button("🗑️", key=f"del_col_{col}_{table_name}", help=f"Delete column {col}"):
                        # Delete column
                        st.session_state[f"edited_df_{table_name}"].drop(columns=[col], inplace=True)
                        df.drop(columns=[col], inplace=True)
                        # Ensure editing state and expander state are maintained
                        st.session_state[f"editing_table_{table_name}"] = True
                        st.session_state[f"expander_open_{table_name}"] = True
                        st.rerun()
                else:
                    st.write("")  # Empty space for alignment
    
    # Add column at the end
    with header_cols[-1]:
        if st.button("➕", key=f"add_col_end_{table_name}", help="Add column at the end"):
            new_col_name = f"Column{len(edited_df.columns) + 1}"
            st.session_state[f"edited_df_{table_name}"][new_col_name] = ""
            df[new_col_name] = ""
            # Ensure editing state and expander state are maintained
            st.session_state[f"editing_table_{table_name}"] = True
            st.session_state[f"expander_open_{table_name}"] = True
            st.rerun()
    
    # Data rows with Excel-like editing (paginated)
    for display_row_num, (row_idx, row) in enumerate(current_page_data.iterrows()):
        row_cols = st.columns(col_widths)
        
        # Calculate actual row number for display (accounting for pagination)
        if total_rows > rows_per_page:
            actual_row_num = (current_page - 1) * rows_per_page + display_row_num + 1
        else:
            actual_row_num = display_row_num + 1
        
        # Row number (always sequential: 1, 2, 3, 4...)
        with row_cols[0]:
            st.markdown(f"**{actual_row_num}**")
        
        # Data cells
        for col_idx, col in enumerate(edited_df.columns):
            with row_cols[col_idx + 1]:  # +1 because first column is row numbers
                # Editable cell - simplified display
                cell_value = str(row[col]) if pd.notna(row[col]) else ""
                new_value = st.text_input(
                    f"Edit {col}",
                    value=cell_value,
                    key=f"cell_{row_idx}_{col}_{table_name}",
                    help=f"Edit: {cell_value}",
                    label_visibility="collapsed"
                )
                
                # Update the dataframe
                st.session_state[f"edited_df_{table_name}"].at[row_idx, col] = new_value
        
        # Row action buttons (add and delete)
        with row_cols[-1]:
            col1, col2 = st.columns([1, 1])
            
            with col1:
                if st.button("➕", key=f"add_row_after_{row_idx}_{table_name}", help=f"Add row after row {actual_row_num}"):
                    # Add new row after current row
                    new_row = pd.Series([''] * len(edited_df.columns), index=edited_df.columns)
                    st.session_state[f"edited_df_{table_name}"] = pd.concat([
                        edited_df.iloc[:row_idx + 1],
                        new_row.to_frame().T,
                        edited_df.iloc[row_idx + 1:]
                    ], ignore_index=True)
                    # Ensure editing state is maintained
                    st.session_state[f"editing_table_{table_name}"] = True
                    st.rerun()
            
            with col2:
                if st.button("🗑️", key=f"del_row_{row_idx}_{table_name}", help=f"Delete row {actual_row_num}"):
                    # Delete current row
                    st.session_state[f"edited_df_{table_name}"].drop(index=row_idx, inplace=True)
                    st.session_state[f"edited_df_{table_name}"].reset_index(drop=True, inplace=True)
                    # Ensure editing state is maintained
                    st.session_state[f"editing_table_{table_name}"] = True
                    st.rerun()
    
    # Close horizontal scroll container if needed
    if len(edited_df.columns) > 6:
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("---")
        
    # Add row at the end
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        if st.button("➕ Add Row at Bottom", key=f"add_row_bottom_{table_name}"):
            new_row = pd.Series([''] * len(edited_df.columns), index=edited_df.columns)
            st.session_state[f"edited_df_{table_name}"] = pd.concat([edited_df, new_row.to_frame().T], ignore_index=True)
            # Ensure editing state is maintained
            st.session_state[f"editing_table_{table_name}"] = True
            st.rerun()
        
        # Action buttons
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 1, 1])
        
        with col1:
            if st.button("✅ Finish Editing", key=f"finish_edit_bottom_{table_name}", type="primary"):
                # Save changes before finishing
                non_empty_rows = st.session_state[f"edited_df_{table_name}"][~(st.session_state[f"edited_df_{table_name}"] == '').all(axis=1)]
                if not non_empty_rows.empty:
                    # Save to appropriate location
                    if table_name in st.session_state.loaded_tables:
                        st.session_state.loaded_tables[table_name] = non_empty_rows
                    elif table_name == 'main':
                        st.session_state.data_processor.data = non_empty_rows
                        st.session_state.data_processor._register_data_in_duckdb()
                    
                    # Clear the edited_df from session state
                    if f"edited_df_{table_name}" in st.session_state:
                        del st.session_state[f"edited_df_{table_name}"]
                    
                    st.success("✅ Changes saved and editing finished!")
                else:
                    st.warning("⚠️ No data to save.")
                
                # Exit editing mode but keep expander open
                st.session_state[f"editing_table_{table_name}"] = False
                st.session_state[f"expander_open_{table_name}"] = True
                st.rerun()
        
        with col2:
            if st.button("🔄 Reset", key=f"reset_data_{table_name}"):
                # Reset to original data
                st.session_state[f"edited_df_{table_name}"] = df.copy()
                st.session_state[f"expander_open_{table_name}"] = True
                st.rerun()
        
        with col3:
            if st.button("❌ Cancel", key=f"cancel_edit_{table_name}"):
                # Discard changes and exit editing mode
                if f"edited_df_{table_name}" in st.session_state:
                    del st.session_state[f"edited_df_{table_name}"]
                st.session_state[f"editing_table_{table_name}"] = False
                st.session_state[f"expander_open_{table_name}"] = True
                st.rerun()

def data_load_tab():
    """Data loading tab with AI model initialization first"""
    st.header("📁 Data Load & AI Setup")
    
    # AI Model Configuration - REQUIRED FIRST STEP
    st.subheader("🤖 Step 1: AI Model Setup (Required)")
    
    # Check if model is already initialized
    model_initialized = hasattr(st.session_state, 'llm_agent') and st.session_state.llm_agent.is_initialized()
    
    if model_initialized:
        st.success(f"✅ AI Model Ready: {st.session_state.llm_agent.model_name}")
        st.markdown("---")
    else:
        st.warning("⚠️ **AI Model Required**: Please initialize an AI model before loading data.")
        
        if OLLAMA_AVAILABLE:
            available_models = get_available_models()
            missing_models = get_missing_models()
            
            if available_models:
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    model_name = st.selectbox(
                        "Select AI Model:",
                        available_models,
                        help="Choose from your available Ollama models",
                        key="data_load_model_selector"
                    )
                
                with col2:
                    if st.button("🔄 Initialize Model", key="data_load_init_model_btn", type="primary"):
                        with st.spinner("Initializing AI model..."):
                            try:
                                success = st.session_state.llm_agent.initialize_model(model_name)
                                if success:
                                    st.success(f"✅ {model_name} initialized!")
                                    st.rerun()
                                else:
                                    st.error(f"❌ Failed to initialize {model_name}")
                            except Exception as e:
                                st.error(f"❌ Error: {str(e)}")
                
                # Show available models info
                st.info(f"📋 **Available Models ({len(available_models)}):** {', '.join(available_models)}")
                
                # Show missing recommended models and download commands
                if missing_models:
                    st.markdown("---")
                    st.subheader("📥 Download Additional Models")
                    st.write("**Recommended models not yet downloaded:**")
                    
                    # Show missing models in a nice format
                    for i, model in enumerate(missing_models[:5], 1):  # Show first 5
                        st.write(f"{i}. `{model}`")
                    
                    if len(missing_models) > 5:
                        st.write(f"... and {len(missing_models) - 5} more")
                    
                    # Generate download commands
                    download_commands = generate_download_commands(missing_models)
                    
                    with st.expander("🔧 Download Commands", expanded=False):
                        st.write("**Copy and run these commands in your terminal:**")
                        st.code("\n".join(download_commands), language="bash")
                        
                        # Individual model download buttons
                        st.write("**Or download individual models:**")
                        cols = st.columns(min(3, len(missing_models[:6])))
                        for i, model in enumerate(missing_models[:6]):
                            with cols[i % 3]:
                                if st.button(f"📥 {model.split(':')[0]}", key=f"download_{model}"):
                                    st.info(f"Run: `ollama pull {model}`")
            else:
                st.warning("⚠️ No AI models found. Please download a model to continue.")
                
                # Model download section with popular models
                st.markdown("### 📥 Download AI Model")
                
                # Get recommended models and show download commands
                recommended_models = get_recommended_models()
                download_commands = generate_download_commands(recommended_models)
                
                st.write("**Popular models to get started:**")
                for i, model in enumerate(recommended_models[:5], 1):
                    st.write(f"{i}. `{model}`")
                
                with st.expander("🔧 Download Commands", expanded=True):
                    st.write("**Copy and run these commands in your terminal:**")
                    st.code("\n".join(download_commands), language="bash")
                
                # Quick start section
                st.markdown("### 🚀 Quick Start")
                st.write("1. **Start Ollama service:** `ollama serve`")
                st.write("2. **Download a model:** `ollama pull llama3:8b-instruct`")
                st.write("3. **Refresh this page** to see available models")
                
                # Popular models section
                st.markdown("**Popular Models:**")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if st.button("📥 Download Llama2", key="download_llama2_btn"):
                        download_model("llama2")
                
                with col2:
                    if st.button("📥 Download Mistral", key="download_mistral_btn"):
                        download_model("mistral")
                
                with col3:
                    if st.button("📥 Download CodeLlama", key="download_codellama_btn"):
                        download_model("codellama")
                
                # Custom model download
                st.markdown("**Or download a custom model:**")
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    model_name = st.text_input(
                        "Model Name:",
                        placeholder="e.g., llama2:7b, mistral:7b, codellama:7b",
                        help="Enter the name of the model you want to download",
                        key="download_model_name"
                    )
                
                with col2:
                    if st.button("📥 Download", key="download_custom_model_btn"):
                        if model_name:
                            download_model(model_name)
                        else:
                            st.warning("⚠️ Please enter a model name.")
                
                # Help section
                st.markdown("---")
                st.markdown("### 💡 Need Help?")
                st.markdown("""
                **To download models manually:**
                1. Open terminal/command prompt
                2. Run: `ollama pull llama2`
                3. Refresh this page
                
                **Popular model sizes:**
                - `llama2` (3.8GB) - Good for general use
                - `llama2:7b` (3.8GB) - Smaller, faster
                - `mistral` (4.1GB) - Great for analysis
                - `codellama` (3.8GB) - Good for SQL generation
                """)
        else:
            st.error("❌ Ollama not available. Please install Ollama first.")
            st.markdown("""
            **Installation Guide:**
            1. Visit: https://ollama.ai/download
            2. Download and install Ollama
            3. Restart this application
            """)
        
        # Disable data loading if model not initialized
        st.markdown("---")
        st.markdown('<div class="status-warning">⚠️ <strong>Data loading is disabled</strong> until AI model is initialized.</div>', unsafe_allow_html=True)
        return  # Exit early if model not initialized
    
    # Data Loading Section - Only shown if model is initialized
    st.subheader("📊 Step 2: Load Your Data")
    
    # File upload section
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("📤 Upload Data Files")
    uploaded_file = st.file_uploader(
            "Choose a data file",
            type=['xlsx', 'xls', 'csv', 'xlsm', 'xlsb', 'odf', 'ods', 'odt'],
            help="Upload your business data in Excel, CSV, or other supported formats",
            key="main_upload"
    )
    
    if uploaded_file is not None:
        # Check if this file has already been processed to avoid reloading on rerun
        file_processed_key = f"file_processed_{uploaded_file.name}_{uploaded_file.size}"
        
        if file_processed_key not in st.session_state:
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension in ['xlsx', 'xls', 'xlsm', 'xlsb']:
                # Multi-sheet Excel file
                with st.spinner("Loading all sheets..."):
                    try:
                        loaded_tables = load_multi_sheet_excel(uploaded_file)
                        if loaded_tables:
                            # Merge Excel data with existing tables instead of replacing
                            if 'loaded_tables' not in st.session_state:
                                st.session_state.loaded_tables = {}
                            
                            # Add Excel tables with a prefix to avoid conflicts
                            # Sanitize the file name for SQL compatibility
                            file_name_without_ext = uploaded_file.name.split('.')[0]
                            sanitized_file_name = sanitize_name(file_name_without_ext)
                            excel_prefix = f"excel_{sanitized_file_name}_"
                            
                            # Show info if file name was sanitized
                            if sanitized_file_name != file_name_without_ext:
                                st.info(f"📝 File name '{file_name_without_ext}' sanitized to '{sanitized_file_name}' for SQL compatibility")
                            tables_added = 0
                            tables_updated = 0
                            
                            for table_name, table_data in loaded_tables.items():
                                # Table name is already sanitized by load_multi_sheet_excel function
                                prefixed_name = f"{excel_prefix}{table_name}"
                                
                                if prefixed_name in st.session_state.loaded_tables:
                                    tables_updated += 1
                                else:
                                    tables_added += 1
                                st.session_state.loaded_tables[prefixed_name] = table_data
                            
                            if tables_updated > 0:
                                st.success(f"✅ Excel file reloaded! {tables_updated} sheets updated, {tables_added} new sheets added. Total tables: {len(st.session_state.loaded_tables)}")
                            else:
                                st.success(f"✅ Loaded {len(loaded_tables)} sheets successfully! Total tables: {len(st.session_state.loaded_tables)}")
                            
                            # Show loaded sheets
                            for sheet_name, df in loaded_tables.items():
                                st.write(f"**{sheet_name}**: {df.shape[0]} rows × {df.shape[1]} columns")
                        else:
                            st.warning("⚠️ No data found in the Excel file")
                    except Exception as e:
                        st.error(f"❌ Error loading Excel file: {str(e)}")
            else:
                # Single file (CSV, etc.)
                try:
                    with st.spinner("Processing file..."):
                        st.session_state.data_processor.load_excel_data(uploaded_file)
                    st.success("✅ Data loaded successfully!")
                except Exception as e:
                    st.error(f"❌ Error loading data: {str(e)}")
            
            # Mark this file as processed
            st.session_state[file_processed_key] = True
    
    # Add Clear All Data button
    if st.session_state.loaded_tables or st.session_state.data_processor.has_data():
        st.markdown("---")
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("🗑️ Clear All Data", key="clear_all_data", type="secondary"):
                # Clear all loaded tables
                st.session_state.loaded_tables = {}
                # Clear data processor
                st.session_state.data_processor.clear_data()
                # Clear file uploader by resetting the session state
                if 'main_upload' in st.session_state:
                    del st.session_state['main_upload']
                # Clear all table-related states
                keys_to_remove = [key for key in st.session_state.keys() if any(
                    key.startswith(prefix) for prefix in ['editing_table_', 'edited_df_', 'renaming_', 'expander_open_', 'file_processed_']
                )]
                for key in keys_to_remove:
                    del st.session_state[key]
                st.success("✅ All data and file uploader cleared successfully!")
                st.rerun()
        
        with col2:
            if st.button("🔄 Refresh Page", key="refresh_page", type="secondary"):
                st.rerun()
    
    with col2:
        st.subheader("📊 Sample Data")
        st.markdown("Try the application with sample business data")
        
        if st.button("📊 Load Sample Data", key="sample_data"):
            with st.spinner("Loading sample data..."):
                try:
                    # Load sample data directly without using data processor to avoid "main" table
                    import os
                    demo_data_path = os.path.join(os.path.dirname(__file__), 'demo_data')
                    
                    loaded_tables = {}
                    
                    # Load sales.xlsx
                    sales_file = os.path.join(demo_data_path, 'sales.xlsx')
                    if os.path.exists(sales_file):
                        sales_df = pd.read_excel(sales_file, sheet_name='Sheet1')
                        loaded_tables['sales'] = sales_df
                    
                    # Load web_traffic.xlsx
                    web_traffic_file = os.path.join(demo_data_path, 'web_traffic.xlsx')
                    if os.path.exists(web_traffic_file):
                        web_traffic_df = pd.read_excel(web_traffic_file, sheet_name='Sheet1')
                        loaded_tables['web_traffic'] = web_traffic_df
                    
                    if loaded_tables:
                        # Merge sample data with existing custom tables instead of replacing
                        if 'loaded_tables' not in st.session_state:
                            st.session_state.loaded_tables = {}
                        
                        # Add sample data tables with a prefix to avoid conflicts
                        sample_prefix = "sample_"
                        tables_added = 0
                        tables_updated = 0
                        
                        for table_name, table_data in loaded_tables.items():
                            # Sanitize the table name for SQL compatibility
                            sanitized_table_name = sanitize_name(table_name)
                            prefixed_name = f"{sample_prefix}{sanitized_table_name}"
                            
                            if prefixed_name in st.session_state.loaded_tables:
                                tables_updated += 1
                            else:
                                tables_added += 1
                            st.session_state.loaded_tables[prefixed_name] = table_data
                        
                        if tables_updated > 0:
                            st.success(f"✅ Sample data reloaded! {tables_updated} tables updated, {tables_added} new tables added. Total tables: {len(st.session_state.loaded_tables)}")
                        else:
                            st.success(f"✅ Sample data loaded! {tables_added} sample tables added. Total tables: {len(st.session_state.loaded_tables)}")
                    else:
                        st.warning("⚠️ No sample data files found in demo_data folder")
                except Exception as e:
                    st.error(f"❌ Error loading sample data: {str(e)}")
    
    # Custom table creation
    st.markdown("---")
    st.subheader("🛠️ Create Custom Table")
    
    with st.expander("➕ Enter Data", expanded=True):
        # Single row interface like PowerBI
        col1, col2 = st.columns([2, 1])
        
        with col1:
            table_name = st.text_input(
                "Table Name:",
                placeholder="e.g., sales_data, customer_info",
                help="Enter a valid SQL table name",
                key="powerbi_table_name"
            )
        
        with col2:
            if st.button("🆕 Create Table", key="powerbi_create_btn", type="primary"):
                if table_name:
                    # Auto-sanitize table name for SQL compatibility
                    original_name = table_name
                    sanitized_table_name = sanitize_name(table_name)
                    
                    # Validate sanitized table name
                    is_valid, message = validate_table_name(sanitized_table_name)
                    if not is_valid:
                        st.error(f"❌ {message}")
                    else:
                        # Use sanitized name
                        table_name = sanitized_table_name
                        
                        # Show info if name was changed
                        if sanitized_table_name != original_name:
                            st.info(f"📝 Table name sanitized from '{original_name}' to '{sanitized_table_name}' for SQL compatibility")
                        # Check for duplicates
                        all_tables = {}
                        if hasattr(st.session_state, 'loaded_tables'):
                            all_tables.update(st.session_state.loaded_tables)
                        if hasattr(st.session_state, 'data_processor') and st.session_state.data_processor.has_data():
                            all_tables['main'] = st.session_state.data_processor.get_data()
                        
                        is_unique, dup_message = check_duplicate_table_name(table_name, all_tables)
                        if not is_unique:
                            st.error(f"❌ {dup_message}")
                        else:
                            # Create empty table with default columns
                            data = {
                                'Column1': pd.Series(dtype='object'),
                                'Column2': pd.Series(dtype='object'),
                                'Column3': pd.Series(dtype='object')
                            }
                            
                            new_df = pd.DataFrame(data)
                            
                            # Initialize loaded_tables if not exists
                            if 'loaded_tables' not in st.session_state:
                                st.session_state.loaded_tables = {}
                            
                            st.session_state.loaded_tables[table_name] = new_df
                            st.session_state.current_table = table_name
                            st.session_state[f"editing_table_{table_name}"] = True
                            
                            st.success(f"✅ Table '{table_name}' created! You can now edit the data directly.")
                            st.rerun()
                else:
                    st.warning("⚠️ Please enter a table name.")
        
        # Show instructions
        st.info("💡 Create a table with default columns, then edit the data directly in the table below.")
    
    # Show loaded data summary
    if st.session_state.loaded_tables or st.session_state.data_processor.has_data():
        st.markdown("---")
        st.subheader("📋 Table Management")
        
        # Display all tables
        all_tables = {}
        if hasattr(st.session_state, 'loaded_tables'):
            all_tables.update(st.session_state.loaded_tables)
        if hasattr(st.session_state, 'data_processor') and st.session_state.data_processor.has_data():
            all_tables['main'] = st.session_state.data_processor.get_data()
        
        if all_tables:
            for table_name, df in all_tables.items():
                # Initialize expander state if not exists
                if f"expander_open_{table_name}" not in st.session_state:
                    st.session_state[f"expander_open_{table_name}"] = False
                
                with st.expander(f"📊 {table_name} ({df.shape[0]} rows × {df.shape[1]} columns)", 
                               expanded=st.session_state[f"expander_open_{table_name}"]):
                    col1, col2, col3, col4, col5 = st.columns([2, 1, 1, 1, 1])
                    
                    with col1:
                        st.markdown(f"**Columns:** {', '.join(df.columns.tolist())}")
                    
                    with col2:
                        st.markdown(f"**Rows:** {df.shape[0]}")
                    
                    with col3:
                        if st.button(f"🏷️ Rename", key=f"rename_{table_name}"):
                            st.session_state[f"renaming_{table_name}"] = True
                            st.rerun()
                    
                    with col4:
                        if st.button(f"✏️ Edit", key=f"edit_{table_name}"):
                            st.session_state[f"editing_table_{table_name}"] = True
                            st.session_state[f"expander_open_{table_name}"] = True
                            st.rerun()
                    
                    with col5:
                        if st.button(f"🗑️ Delete", key=f"delete_{table_name}"):
                            # Delete from loaded_tables if it exists there
                            if table_name in st.session_state.loaded_tables:
                                del st.session_state.loaded_tables[table_name]
                                st.success(f"✅ Table '{table_name}' deleted successfully!")
                                
                                # Clear any editing states for this table
                                if f"editing_table_{table_name}" in st.session_state:
                                    del st.session_state[f"editing_table_{table_name}"]
                                if f"edited_df_{table_name}" in st.session_state:
                                    del st.session_state[f"edited_df_{table_name}"]
                                if f"renaming_{table_name}" in st.session_state:
                                    del st.session_state[f"renaming_{table_name}"]
                                
                                # Don't clear file uploader state - let user manage it manually
                                # This allows users to delete tables without losing their uploaded file
                                    
                            # Also clear from data_processor if it's the main table
                            elif table_name == 'main':
                                st.session_state.data_processor.clear_data()
                                st.success(f"✅ Main table cleared successfully!")
                                
                                # Clear file uploader if it exists
                                if 'main_upload' in st.session_state:
                                    del st.session_state['main_upload']
                                    
                            # Handle any other cases
                            else:
                                # Try to clear from data_processor as well
                                st.session_state.data_processor.clear_data()
                                st.success(f"✅ Table '{table_name}' deleted successfully!")
                                
                                # Clear file uploader if it exists
                                if 'main_upload' in st.session_state:
                                    del st.session_state['main_upload']
                            st.rerun()
                    
                    # Rename functionality
                    if st.session_state.get(f"renaming_{table_name}", False):
                        new_name = st.text_input(f"New name for '{table_name}':", value=table_name)
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("✅ Confirm", key=f"confirm_rename_{table_name}"):
                                if new_name and new_name != table_name:
                                    is_valid, message = validate_table_name(new_name)
                                    if is_valid:
                                        # Create a copy of all_tables without the current table being renamed
                                        tables_to_check = {k: v for k, v in all_tables.items() if k != table_name}
                                        is_unique, dup_message = check_duplicate_table_name(new_name, tables_to_check)
                                        if is_unique:
                                            # Get the table data first
                                            table_data = None
                                            
                                            # Check where the table exists and get its data
                                            if table_name in st.session_state.loaded_tables:
                                                table_data = st.session_state.loaded_tables.pop(table_name)
                                            elif table_name == 'main' and st.session_state.data_processor.has_data():
                                                table_data = st.session_state.data_processor.get_data()
                                                st.session_state.data_processor.clear_data()
                                            
                                            if table_data is not None:
                                                # Add the table with new name to loaded_tables
                                                st.session_state.loaded_tables[new_name] = table_data
                                                
                                                # Update current table reference
                                                if st.session_state.get('current_table') == table_name:
                                                    st.session_state.current_table = new_name
                                                
                                                # Clear any editing states for the old table name
                                                if f"editing_table_{table_name}" in st.session_state:
                                                    del st.session_state[f"editing_table_{table_name}"]
                                                if f"edited_df_{table_name}" in st.session_state:
                                                    del st.session_state[f"edited_df_{table_name}"]
                                                
                                                # Clear the renaming state
                                                if f"renaming_{table_name}" in st.session_state:
                                                    del st.session_state[f"renaming_{table_name}"]
                                                
                                                # Also clear any file uploader state that might be tied to this table
                                                # This ensures the file uploader doesn't interfere with table operations
                                                if 'main_upload' in st.session_state and not st.session_state.loaded_tables:
                                                    # Only clear if no other tables exist
                                                    pass  # Keep the file uploader state
                                                
                                                st.success(f"✅ Table renamed from '{table_name}' to '{new_name}'")
                                                st.rerun()
                                            else:
                                                st.error(f"❌ Could not find table '{table_name}' to rename")
                                        else:
                                            st.error(f"❌ {dup_message}")
                                    else:
                                        st.error(f"❌ {message}")
                                else:
                                    st.error("❌ New name cannot be empty or same as current name")
                        with col2:
                            if st.button("❌ Cancel", key=f"cancel_rename_{table_name}"):
                                st.session_state[f"renaming_{table_name}"] = False
                                st.rerun()
    
                    # Check if this table is being edited
                    if st.session_state.get(f"editing_table_{table_name}", False):
                        # Show editor interface instead of data preview
                        st.markdown("**📝 Table Editor:**")
                        show_table_editor_interface(table_name, df)
                    else:
                        # Show data preview
                        st.markdown("**Data Preview:**")
                        safe_dataframe_display(df.head(10), width='stretch')
    

def data_explorer_tab():
    """Data exploration tab with column selection"""
    st.header("🔍 Explore Your Data")
    
    if not st.session_state.loaded_tables and not st.session_state.data_processor.has_data():
        st.markdown('<div class="status-warning">⚠️ No data loaded. Please go to the Data Load tab to upload your data.</div>', unsafe_allow_html=True)
        return
    
    # Table selection - include both loaded tables and main data
    all_available_tables = {}
    if st.session_state.loaded_tables:
        all_available_tables.update(st.session_state.loaded_tables)
    if st.session_state.data_processor.has_data():
        all_available_tables['main'] = st.session_state.data_processor.get_data()
    
    if all_available_tables:
        selected_table = st.selectbox(
            "Select Table to Explore:",
            list(all_available_tables.keys()),
            key="table_selector_explore"
        )
        
        if selected_table:
            current_data = all_available_tables[selected_table]
            st.session_state.current_table = selected_table
    else:
        current_data = None
    
    if current_data is not None and not current_data.empty:
        # Column selection
        st.subheader("📋 Select Columns for Analysis")
        
        all_columns = list(current_data.columns)
        selected_columns = st.multiselect(
            "Choose columns to analyze:",
            all_columns,
            default=all_columns[:5] if len(all_columns) > 5 else all_columns,
            key=f"column_selector_{st.session_state.current_table}"
        )
        
        if selected_columns:
            st.session_state.selected_columns[st.session_state.current_table] = selected_columns
            filtered_data = current_data[selected_columns]
            
            # Data overview
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Selected Columns", len(selected_columns))
            with col2:
                st.metric("Total Rows", len(filtered_data))
            with col3:
                st.metric("Memory Usage", f"{filtered_data.memory_usage(deep=True).sum() / 1024 / 1024:.2f} MB")
            with col4:
                missing_pct = (filtered_data.isnull().sum().sum() / (len(filtered_data) * len(selected_columns))) * 100
                st.metric("Missing Data", f"{missing_pct:.1f}%")
            
            # Data preview
            st.subheader("👀 Data Preview")
            show_rows = st.slider("Rows to display", 5, 100, 10)
            safe_dataframe_display(filtered_data.head(show_rows), width='stretch')
            
            # Data analysis
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📊 Data Types")
                type_counts = filtered_data.dtypes.value_counts()
                type_names = [str(dtype) for dtype in type_counts.index]
                fig = px.pie(values=type_counts.values, names=type_names, title="Column Types")
                st.plotly_chart(fig, width='stretch')
            
            with col2:
                st.subheader("❓ Missing Values")
                missing_data = filtered_data.isnull().sum()
                if missing_data.sum() > 0:
                    fig = px.bar(x=missing_data.index, y=missing_data.values, title="Missing Values per Column")
                    st.plotly_chart(fig, width='stretch')
                else:
                    st.markdown('<div class="status-success">✅ No missing values found!</div>', unsafe_allow_html=True)
            
            # Statistical summary
            st.subheader("📈 Statistical Summary")
            safe_dataframe_display(filtered_data.describe(), width='stretch')
    
    # Navigation to next tab

def edit_relationship_interface(rel_index, current_rel, all_tables):
    """Interface for editing an existing relationship"""
    st.markdown("**Edit Relationship Details**")
    
    # Get available tables and columns
    table_names = list(all_tables.keys())
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Source Table & Column")
        source_table = st.selectbox(
            "Source Table",
            table_names,
            index=table_names.index(current_rel['source_table']) if current_rel['source_table'] in table_names else 0,
            key=f"edit_source_table_{rel_index}"
        )
        
        if source_table in all_tables:
            source_columns = all_tables[source_table].columns.tolist()
            source_column = st.selectbox(
                "Source Column",
                source_columns,
                index=source_columns.index(current_rel['source_column']) if current_rel['source_column'] in source_columns else 0,
                key=f"edit_source_column_{rel_index}"
            )
        else:
            st.warning("Source table not found")
            return
    
    with col2:
        st.subheader("Target Table & Column")
        target_table = st.selectbox(
            "Target Table",
            table_names,
            index=table_names.index(current_rel['target_table']) if current_rel['target_table'] in table_names else 0,
            key=f"edit_target_table_{rel_index}"
        )
        
        if target_table in all_tables:
            target_columns = all_tables[target_table].columns.tolist()
            target_column = st.selectbox(
                "Target Column",
                target_columns,
                index=target_columns.index(current_rel['target_column']) if current_rel['target_column'] in target_columns else 0,
                key=f"edit_target_column_{rel_index}"
            )
        else:
            st.warning("Target table not found")
            return
    
    # Relationship type
    st.subheader("Relationship Type")
    rel_type = st.selectbox(
        "Relationship Type",
        ["One-to-One", "One-to-Many", "Many-to-One", "Many-to-Many"],
        index=["One-to-One", "One-to-Many", "Many-to-One", "Many-to-Many"].index(current_rel.get('type', 'One-to-Many')),
        key=f"edit_rel_type_{rel_index}"
    )
    
    # Action buttons
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        if st.button("💾 Save Changes", key=f"save_edit_{rel_index}", type="primary"):
            # Update the relationship
            st.session_state.relationships[rel_index] = {
                'source_table': source_table,
                'source_column': source_column,
                'target_table': target_table,
                'target_column': target_column,
                'type': rel_type
            }
            
            # Clear editing state
            if f"editing_rel_{rel_index}" in st.session_state:
                del st.session_state[f"editing_rel_{rel_index}"]
            
            st.success("✅ Relationship updated successfully!")
            st.rerun()
    
    with col2:
        if st.button("❌ Cancel", key=f"cancel_edit_{rel_index}"):
            # Clear editing state
            if f"editing_rel_{rel_index}" in st.session_state:
                del st.session_state[f"editing_rel_{rel_index}"]
            st.rerun()
    
    with col3:
        if st.button("🗑️ Delete", key=f"delete_edit_{rel_index}"):
            # Delete the relationship
            st.session_state.relationships.pop(rel_index)
            
            # Clear editing state
            if f"editing_rel_{rel_index}" in st.session_state:
                del st.session_state[f"editing_rel_{rel_index}"]
            
            st.success("✅ Relationship deleted!")
            st.rerun()
    

def relationship_builder_tab():
    """Relationship builder tab with ER diagram"""
    st.header("🔗 Build Table Relationships")
    
    if not st.session_state.loaded_tables and not st.session_state.data_processor.has_data():
        st.markdown('<div class="status-warning">⚠️ No data loaded. Please load data first.</div>', unsafe_allow_html=True)
        return
    
    # Get all available tables
    all_tables = {}
    if st.session_state.loaded_tables:
        all_tables.update(st.session_state.loaded_tables)
    if st.session_state.data_processor.has_data():
        all_tables.update(st.session_state.data_processor.get_loaded_tables())
    
    if not all_tables:
        st.markdown('<div class="status-warning">⚠️ No tables available for relationship building.</div>', unsafe_allow_html=True)
        return
    
    # Display table metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Tables", len(all_tables))
    
    with col2:
        total_cols = sum(df.shape[1] for df in all_tables.values())
        st.metric("Total Columns", total_cols)
    
    with col3:
        total_rows = sum(df.shape[0] for df in all_tables.values())
        st.metric("Total Rows", f"{total_rows:,}")
    
    with col4:
        relationships_count = len(st.session_state.get('relationships', []))
        st.metric("Relationships", relationships_count)
    
    # Main relationship builder interface - single column layout
    st.subheader("🔗 Create Relationship")
    
    # AI-powered relationship detection
    if st.button("🤖 Auto-Detect Relationships", key="ai_detect_relationships_btn", type="primary", use_container_width=True):
        if hasattr(st.session_state, 'llm_agent') and st.session_state.llm_agent.is_initialized():
            # Reset detection state to allow new detection
            st.session_state[f"ai_detecting_relationships"] = True
            st.session_state[f"ai_detection_started"] = False
            st.session_state[f"ai_detection_completed"] = False
            st.session_state[f"ai_detection_timestamp"] = st.session_state.get("_last_rerun_timestamp", 0)
            st.rerun()
        else:
            st.warning("⚠️ Please initialize an AI model in the Data Load tab first.")
    
    # Show AI detection progress and results
    if st.session_state.get(f"ai_detecting_relationships", False) and not st.session_state.get(f"ai_detection_completed", False):
        # Check if this is the first time we're running detection
        if not st.session_state.get(f"ai_detection_started", False):
            st.session_state[f"ai_detection_started"] = True
            st.session_state[f"ai_detection_completed"] = False
            
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Stop button
        col_stop1, col_stop2 = st.columns([1, 4])
        with col_stop1:
            if st.button("⏹️ Stop AI Detection", key="stop_ai_detection_btn"):
                st.session_state[f"ai_detecting_relationships"] = False
                st.session_state[f"ai_detection_started"] = False
                st.session_state[f"ai_detection_completed"] = True
                st.rerun()
        
        with col_stop2:
            st.write("")  # Empty space for alignment
        
        # Simulate AI detection process
        try:
            # Step 1: Analyzing table structure
            progress_bar.progress(20)
            status_text.text("🔍 Analyzing table structure...")
            
            # Step 2: Detecting relationships
            progress_bar.progress(60)
            status_text.text("🤖 AI is detecting relationships...")
            
            # Call the actual AI detection
            detected_relationships = st.session_state.llm_agent.detect_relationships(all_tables)
            
            # Step 3: Validating relationships
            progress_bar.progress(90)
            status_text.text("✅ Validating relationships...")
            
            if detected_relationships:
                # Filter out invalid relationships
                valid_relationships = []
                invalid_relationships = 0
                
                for rel in detected_relationships:
                    # Check if all required fields exist and are valid
                    if (all(field in rel for field in ['source_table', 'target_table', 'source_column', 'target_column']) and
                        rel['source_table'] in all_tables and
                        rel['target_table'] in all_tables and
                        rel['source_column'] in all_tables[rel['source_table']].columns and
                        rel['target_column'] in all_tables[rel['target_table']].columns):
                        valid_relationships.append(rel)
                    else:
                        invalid_relationships += 1
                
                if valid_relationships:
                    # Add valid relationships to session state
                    if 'relationships' not in st.session_state:
                        st.session_state.relationships = []
                    
                    # Add only new relationships (avoid duplicates)
                    new_relationships = 0
                    for rel in valid_relationships:
                        # Check if this relationship already exists
                        exists = any(
                            existing['source_table'] == rel['source_table'] and
                            existing['target_table'] == rel['target_table'] and
                            existing['source_column'] == rel['source_column'] and
                            existing['target_column'] == rel['target_column']
                            for existing in st.session_state.relationships
                        )
                        
                        if not exists:
                            # Convert AI format to our format
                            relationship = {
                                'source_table': rel['source_table'],
                                'target_table': rel['target_table'],
                                'source_column': rel['source_column'],
                                'target_column': rel['target_column'],
                                'type': rel.get('relationship_type', 'One-to-Many'),
                                'confidence': rel.get('confidence', 'Medium'),
                                'reasoning': rel.get('reasoning', 'AI detected'),
                                'detection_method': rel.get('detection_method', 'ai')
                            }
                            st.session_state.relationships.append(relationship)
                            new_relationships += 1
                    
                    progress_bar.progress(100)
                    status_text.text("✅ Detection completed!")
                    
                    if new_relationships > 0:
                        st.markdown("**🔍 Validated Relationships Added:**")
                        for i, rel in enumerate(detected_relationships[-new_relationships:]):
                            if rel.get('validated', False):
                                st.markdown(f"• **{rel['source_table']}.{rel['source_column']}** → **{rel['target_table']}.{rel['target_column']}** ({rel.get('relationship_type', 'One-to-Many')})")
                        st.info("💡 These relationships are now available for editing in the 'Current Relationships' section below!")
                else:
                    status_text.text("🤖 No valid relationships detected.")
                    st.warning("🤖 AI couldn't detect any valid relationships. All suggested relationships had non-existent columns.")
                    if invalid_relationships > 0:
                        st.info(f"ℹ️ {invalid_relationships} relationships were suggested but filtered out due to invalid columns.")
                
                st.session_state[f"ai_detecting_relationships"] = False
                st.session_state[f"ai_detection_completed"] = True
                st.rerun()
            else:
                progress_bar.progress(100)
                status_text.text("🤖 No relationships detected.")
                st.warning("🤖 AI couldn't detect any obvious relationships.")
                # Completely stop detection - don't show completion UI
                st.session_state[f"ai_detecting_relationships"] = False
                st.session_state[f"ai_detection_started"] = False
                st.session_state[f"ai_detection_completed"] = True  # Mark as completed even with no results
                st.rerun()
                
        except Exception as e:
            progress_bar.progress(100)
            status_text.text("❌ Detection failed.")
            st.error(f"❌ Error in AI relationship detection: {str(e)}")
            print(f"AI relationship detection error: {str(e)}")
            # Completely stop detection - don't show completion UI
            st.session_state[f"ai_detecting_relationships"] = False
            st.session_state[f"ai_detection_started"] = False
            st.session_state[f"ai_detection_completed"] = True  # Mark as completed even with error
    
    # Manual relationship creation
    st.markdown("---")
    st.markdown("**Manual Relationship Creation**")
    
    # Table selection
    table1 = st.selectbox("Source Table", list(all_tables.keys()), key="rel_table1")
    table2 = st.selectbox("Target Table", list(all_tables.keys()), key="rel_table2")
    
    if table1 and table2 and table1 != table2:
        col1_col, col2_col = st.columns(2)
        
        with col1_col:
            source_col = st.selectbox("Source Column", all_tables[table1].columns.tolist(), key="rel_source_col")
        
        with col2_col:
            target_col = st.selectbox("Target Column", all_tables[table2].columns.tolist(), key="rel_target_col")
        
        rel_type = st.selectbox("Relationship Type", ["One-to-One", "One-to-Many", "Many-to-One", "Many-to-Many"], key="rel_type")
        
        if st.button("➕ Add Relationship", key="add_manual_rel"):
            if 'relationships' not in st.session_state:
                st.session_state.relationships = []
            
            new_rel = {
                'source_table': table1,
                'source_column': source_col,
                'target_table': table2,
                'target_column': target_col,
                'type': rel_type
            }
            
            # Check for duplicates
            if new_rel not in st.session_state.relationships:
                st.session_state.relationships.append(new_rel)
                st.success(f"✅ Relationship added: {table1}.{source_col} → {table2}.{target_col} ({rel_type})")
                st.rerun()
            else:
                st.warning("⚠️ This relationship already exists.")
    
    st.markdown("---")
    
    # Current Relationships section - moved below and improved
    st.subheader("📊 Current Relationships")
    
    # Initialize relationships if not exists
    if 'relationships' not in st.session_state:
        st.session_state.relationships = []
    
    # Relationship count and search
    rel_count = len(st.session_state.relationships)
    if rel_count > 0:
        # Count AI detected relationships
        ai_detected_count = sum(1 for rel in st.session_state.relationships if 'confidence' in rel)
        manual_count = rel_count - ai_detected_count
        
        # Improved relationship display with better layout
        col1, col2 = st.columns([3, 1])
        
        with col1:
            st.markdown(f"**Total: {rel_count} relationship{'s' if rel_count != 1 else ''}**")
            if ai_detected_count > 0:
                st.markdown(f"• 🤖 AI Detected: {ai_detected_count}")
            if manual_count > 0:
                st.markdown(f"• ✋ Manual: {manual_count}")
        
        with col2:
            # Search/filter functionality
            search_term = st.text_input(
                "🔍 Search relationships",
                key="relationship_search",
                placeholder="Search by table or column name...",
                help="Filter relationships by table or column names"
            )
        
        # Filter relationships based on search
        if search_term:
            filtered_relationships = [
                rel for rel in st.session_state.relationships
                if (search_term.lower() in rel['source_table'].lower() or
                    search_term.lower() in rel['target_table'].lower() or
                    search_term.lower() in rel['source_column'].lower() or
                    search_term.lower() in rel['target_column'].lower())
            ]
            st.info(f"Found {len(filtered_relationships)} relationship{'s' if len(filtered_relationships) != 1 else ''} matching '{search_term}'")
        else:
            filtered_relationships = st.session_state.relationships
    else:
        filtered_relationships = []
    
    # Display relationships in a more legible format
    if filtered_relationships:
        st.markdown("---")
        
        for i, rel in enumerate(filtered_relationships):
            # Get the actual index in the original relationships list
            original_index = st.session_state.relationships.index(rel)
            
            # Check if this is a recently detected relationship (has confidence field)
            is_ai_detected = 'confidence' in rel
            border_color = "#28a745" if is_ai_detected else "#e0e0e0"
            bg_color = "#d4edda" if is_ai_detected else "#f8f9fa"
            ai_badge = "🤖 AI Detected" if is_ai_detected else "✋ Manual"
            
            # Create a more legible relationship card
            with st.container():
                st.markdown(f"""
                <div style="border: 2px solid {border_color}; border-radius: 8px; padding: 16px; margin: 12px 0; background-color: {bg_color};">
                    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;">
                        <h4 style="margin: 0; color: #333;">🔗 {rel['source_table']}.{rel['source_column']} → {rel['target_table']}.{rel['target_column']}</h4>
                        <span style="background-color: {'#28a745' if is_ai_detected else '#6c757d'}; color: white; padding: 4px 8px; border-radius: 4px; font-size: 12px; font-weight: bold;">{ai_badge}</span>
                    </div>
                    <div style="display: flex; justify-content: space-between; align-items: center;">
                        <div>
                            <strong>Type:</strong> {rel['type']}<br>
                            <small style="color: #666;">Source: {rel['source_table']} | Target: {rel['target_table']}</small>
                        </div>
                        <div style="display: flex; gap: 8px;">
                            <button onclick="editRelationship({original_index})" style="background: #007bff; color: white; border: none; padding: 6px 12px; border-radius: 4px; cursor: pointer;">✏️ Edit</button>
                            <button onclick="deleteRelationship({original_index})" style="background: #dc3545; color: white; border: none; padding: 6px 12px; border-radius: 4px; cursor: pointer;">🗑️ Delete</button>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            # Action buttons using Streamlit
            col1, col2, col3 = st.columns([1, 1, 8])
            
            with col1:
                if st.button("✏️", key=f"edit_rel_{original_index}", help="Edit this relationship"):
                    st.session_state[f"editing_rel_{original_index}"] = True
                    st.rerun()
            
            with col2:
                if st.button("🗑️", key=f"delete_rel_{original_index}", help="Delete this relationship"):
                    st.session_state.relationships.pop(original_index)
                    st.rerun()
            
            with col3:
                st.write("")  # Empty space for alignment
            
            # Edit interface for this relationship
            if st.session_state.get(f"editing_rel_{original_index}", False):
                with st.expander(f"✏️ Edit Relationship {original_index+1}", expanded=True):
                    edit_relationship_interface(original_index, rel, all_tables)
    else:
        st.info("No relationships defined yet. Use AI detection or manual creation to add relationships.")
    
    # Bulk operations
    if st.session_state.relationships:
        st.markdown("---")
        st.subheader("🔧 Bulk Operations")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔄 Refresh All", key="refresh_all_relationships", help="Refresh all relationships"):
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All", key="clear_all_relationships", help="Delete all relationships"):
                st.session_state.relationships = []
                st.success("✅ All relationships cleared!")
                st.rerun()
        
        with col3:
            if st.button("📊 Export Relationships", key="export_relationships", help="Export relationships to JSON"):
                import json
                relationships_json = json.dumps(st.session_state.relationships, indent=2)
                st.download_button(
                    label="📥 Download Relationships",
                    data=relationships_json,
                    file_name="relationships.json",
                    mime="application/json"
                )
        
        # Import relationships
        st.markdown("**📤 Import Relationships**")
        uploaded_file = st.file_uploader(
            "Upload relationships JSON file",
            type=['json'],
            key="import_relationships_file",
            help="Upload a JSON file containing relationships"
        )
        
        if uploaded_file is not None:
            try:
                import json
                relationships_data = json.load(uploaded_file)
                if isinstance(relationships_data, list):
                    st.session_state.relationships = relationships_data
                    st.success(f"✅ Imported {len(relationships_data)} relationships!")
                    st.rerun()
                else:
                    st.error("❌ Invalid JSON format. Expected a list of relationships.")
            except Exception as e:
                st.error(f"❌ Error importing relationships: {str(e)}")
    
    # ER Diagram section
    st.markdown("---")
    st.subheader("📊 Entity Relationship Diagram")
    
    if st.session_state.relationships:
        if st.button("🔄 Refresh Diagram", key="refresh_diagram_btn"):
            st.rerun()
        
        er_diagram = create_er_diagram(all_tables, st.session_state.relationships)
        if er_diagram:
            st.pyplot(er_diagram, width='stretch')
    else:
        st.info("No relationships defined yet. Create some relationships to see the ER diagram.")
def ai_analysis_tab():
    """AI analysis tab with improved workflow and persistent SQL queries"""
    st.header("🤖 AI-Powered Analysis")
    
    if not st.session_state.loaded_tables and not st.session_state.data_processor.has_data():
        st.markdown('<div class="status-warning">⚠️ No data loaded. Please load data first.</div>', unsafe_allow_html=True)
        return
    
    # Check if model is initialized
    if not hasattr(st.session_state, 'llm_agent') or not st.session_state.llm_agent.is_initialized():
        st.warning("⚠️ Please initialize an AI model in the Data Load tab first.")
        return
    
    # AI Workflow: Combined Question Input and SQL Generation
    st.subheader("💬 Ask Your Question & Generate SQL")
    
    col1, col2 = st.columns([4, 1])
    
    with col1:
        # Determine what value to show in the text area
        display_value = st.session_state.question_to_display if st.session_state.question_to_display else st.session_state.analysis_query
        
        user_query = st.text_area(
            "Enter your analysis question:",
            value=display_value,
            placeholder="e.g., What are the top 5 products by sales? Show me customer trends over time.",
            height=100,
            key="analysis_query"
        )
    
    # Handle question generation flag
    if st.session_state.generate_question_flag:
        with st.spinner("AI is analyzing your data and generating an interesting question..."):
            try:
                # Get all available tables (same logic as Generate SQL)
                all_tables = {}
                if st.session_state.loaded_tables:
                    all_tables.update(st.session_state.loaded_tables)
                if st.session_state.data_processor.has_data():
                    all_tables['main'] = st.session_state.data_processor.get_data()
                
                # Get relationships for question generation (same as Generate SQL)
                relationships = st.session_state.get('relationships', [])
                
                # Generate an interesting analysis question with all table information and relationships
                generated_question = st.session_state.llm_agent.generate_analysis_question_with_tables(all_tables, relationships)
                
                # Store the generated question in separate variables (don't modify analysis_query)
                st.session_state.generated_question = generated_question
                st.session_state.question_to_display = generated_question
                
                # Reset the flag
                st.session_state.generate_question_flag = False
                
                # Show success message and rerun to update the display
                st.success("✅ Generated an interesting analysis question! Click 'Generate SQL' to create the query.")
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
                st.session_state.generate_question_flag = False
    
    # Clear display variable if user has manually edited the text area
    if user_query != st.session_state.question_to_display and st.session_state.question_to_display:
        st.session_state.question_to_display = ""
    
    with col2:
        if st.button("🔍 Generate SQL", type="primary", key="generate_sql_btn"):
            if user_query:
                with st.spinner("AI is analyzing your question and generating SQL..."):
                    try:
                        # Get all available tables
                        all_tables = {}
                        if st.session_state.loaded_tables:
                            all_tables.update(st.session_state.loaded_tables)
                        if st.session_state.data_processor.has_data():
                            all_tables['main'] = st.session_state.data_processor.get_data()
                        
                        # Get relationships for SQL generation
                        relationships = st.session_state.get('relationships', [])
                        
                        # Generate SQL query with all table information and relationships
                        sql_query = st.session_state.llm_agent.generate_sql_query_with_tables(user_query, all_tables, relationships)
                        
                        # Store SQL query
                        st.session_state.ai_analysis_sql_query = sql_query
                        st.session_state.current_user_query = user_query
                        
                        st.success("✅ SQL query generated!")
                        
                    except Exception as e:
                        st.error(f"❌ Error generating SQL: {str(e)}")
        
        with col2:
            if st.button("🎯 Generate Question", key="quick_sql_btn"):
                # Set flag to generate question in next execution
                st.session_state.generate_question_flag = True
                st.rerun()
    
    # Display generated SQL query (persistent and editable)
    if hasattr(st.session_state, 'ai_analysis_sql_query') and st.session_state.ai_analysis_sql_query:
        st.markdown("### 📝 SQL Query:")
        
        # Editable SQL query
        edited_sql = st.text_area(
            "Edit SQL Query:",
            value=st.session_state.ai_analysis_sql_query,
            height=150,
            help="You can edit the generated SQL query or write your own custom query",
            key="ai_analysis_sql_editor"
        )
        
        # Update the stored SQL query if edited
        if edited_sql != st.session_state.ai_analysis_sql_query:
            st.session_state.ai_analysis_sql_query = edited_sql
        
        # AI Workflow: Step 3 - Execute Query
        st.subheader("📊 Step 3: Execute Query")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            if st.button("▶️ Execute Query", type="primary", key="execute_query_ai_btn"):
                with st.spinner("Executing query..."):
                    try:
                        # Get all available tables
                        all_tables = {}
                        if st.session_state.loaded_tables:
                            all_tables.update(st.session_state.loaded_tables)
                        if st.session_state.data_processor.has_data():
                            all_tables['main'] = st.session_state.data_processor.get_data()
                        
                        # Execute SQL query
                        try:
                            # Register all tables in DuckDB
                            conn = duckdb.connect()
                            for table_name, df in all_tables.items():
                                conn.register(table_name, df)
                            
                            # Execute query
                            result_df = conn.execute(st.session_state.ai_analysis_sql_query).fetchdf()
                            
                            if not result_df.empty:
                                st.success("✅ Query executed successfully!")
                                
                                # Store results
                                st.session_state.ai_analysis_query_results = result_df
                                
                            else:
                                st.warning("⚠️ Query returned no results.")
                                
                        except Exception as sql_error:
                            st.error(f"❌ SQL execution error: {str(sql_error)}")
                            st.info("💡 Try rephrasing your question or check the SQL query.")
                        
                    except Exception as e:
                        st.error(f"❌ Error executing query: {str(e)}")
        
        with col2:
            if st.button("🔄 Re-generate SQL", key="regenerate_sql_btn"):
                st.session_state.ai_analysis_sql_query = None
                st.rerun()
    
    # Display persistent query results if available
    if hasattr(st.session_state, 'ai_analysis_query_results') and st.session_state.ai_analysis_query_results is not None:
        st.markdown("---")
        st.subheader("📊 Query Results")
        
        result_df = st.session_state.ai_analysis_query_results
        
        # Display results summary
        st.markdown("### 📋 Results Summary:")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Rows Returned", len(result_df))
        with col2:
            st.metric("Columns", len(result_df.columns))
        with col3:
            st.metric("Memory Usage", f"{result_df.memory_usage(deep=True).sum() / 1024 / 1024:.2f} MB")
        
        # Display results
        st.markdown("### 📊 Data Preview:")
        safe_dataframe_display(result_df, width='stretch')
    
    
    
    # AI Workflow: Step 4 - Visualization (Separate Button)
    if (hasattr(st.session_state, 'ai_analysis_query_results') and 
        st.session_state.ai_analysis_query_results is not None and 
        not st.session_state.ai_analysis_query_results.empty):
        st.subheader("📈 Step 4: Create Visualization")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            if st.button("📊 Generate Visualization", type="primary", key="generate_viz_btn"):
                with st.spinner("Creating visualization..."):
                    try:
                        # Create visualization
                        user_query = st.session_state.get('current_user_query', 'Custom Query Results')
                        viz_fig = st.session_state.visualizer.create_auto_visualization(
                            st.session_state.ai_analysis_query_results, 
                            user_query
                        )
                        if viz_fig:
                            st.session_state.current_visualization = viz_fig
                            st.success("✅ Visualization created!")
                        else:
                            st.warning("⚠️ Could not create visualization for this data.")
                    except Exception as e:
                        st.error(f"❌ Error creating visualization: {str(e)}")
    
    with col2:
            if st.button("🤖✨ Generate AI Insights", type="primary", key="generate_insights_btn"):
                with st.spinner("AI is analyzing the results..."):
                    try:
                        # Get source tables for enhanced analysis
                        source_tables = {}
                        if hasattr(st.session_state, 'loaded_tables'):
                            source_tables.update(st.session_state.loaded_tables)
                        if hasattr(st.session_state, 'data_processor') and st.session_state.data_processor.has_data():
                            source_tables['main'] = st.session_state.data_processor.get_data()
                        
                        # Generate AI analysis with source tables context
                        user_query = st.session_state.get('current_user_query', 'Custom Query Analysis')
                        analysis = st.session_state.llm_agent.analyze_query_results(
                            user_query,
                            st.session_state.ai_analysis_query_results,
                            source_tables
                        )
                        
                        # Store analysis
                        st.session_state.current_ai_analysis = analysis
                        st.session_state.analysis_timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        
                        st.success("✅ AI analysis complete!")
                        
                    except Exception as e:
                        st.error(f"❌ Error generating analysis: {str(e)}")
    
    # Display Visualization
    if hasattr(st.session_state, 'current_visualization') and st.session_state.current_visualization:
        st.markdown("### 📈 Visualization:")
        st.plotly_chart(st.session_state.current_visualization, width='stretch')
    
    # Display AI Analysis
    if hasattr(st.session_state, 'current_ai_analysis') and st.session_state.current_ai_analysis:
        st.markdown("### 🧠 AI Analysis:")
        escaped_analysis = escape_markdown_for_streamlit(st.session_state.current_ai_analysis)
        st.markdown(escaped_analysis)
        
        # Store in insights history
        if 'insights_history' not in st.session_state:
            st.session_state.insights_history = []
        
        insight_entry = {
            'question': st.session_state.get('current_user_query', 'Custom Query Analysis'),
            'sql_query': st.session_state.get('ai_analysis_sql_query', ''),
            'analysis': st.session_state.current_ai_analysis,
            'timestamp': st.session_state.analysis_timestamp,
            'data_shape': st.session_state.ai_analysis_query_results.shape if st.session_state.ai_analysis_query_results is not None else (0, 0)
        }
        
        # Add to history if not already there
        if not any(entry['timestamp'] == insight_entry['timestamp'] for entry in st.session_state.insights_history):
            st.session_state.insights_history.append(insight_entry)
        
        # Export options for current analysis
        st.markdown("---")
        st.subheader("📄 Export Options")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📄 Export Current as PDF", key="export_current_pdf", type="primary"):
                try:
                    if REPORTLAB_AVAILABLE:
                        pdf_data = generate_pdf_report(insight_entry, f"ai_analysis_{insight_entry['timestamp'].replace(':', '-').replace(' ', '_')}.pdf")
                        st.download_button(
                            label="📥 Download PDF",
                            data=pdf_data,
                            file_name=f"ai_analysis_{insight_entry['timestamp'].replace(':', '-').replace(' ', '_')}.pdf",
                            mime="application/pdf"
                        )
                    else:
                        st.error("❌ ReportLab not available. Install with: pip install reportlab")
                except Exception as e:
                    st.error(f"❌ Error creating PDF: {str(e)}")
        
        with col2:
            if st.button("📊 Export All Insights as PDF", key="export_all_pdf", type="secondary"):
                try:
                    if REPORTLAB_AVAILABLE:
                        # Generate bulk PDF with all insights
                        pdf_data = generate_bulk_pdf_report(st.session_state.insights_history)
                        st.download_button(
                            label="📥 Download All Insights PDF",
                            data=pdf_data,
                            file_name=f"all_ai_insights_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf"
                        )
                    else:
                        st.error("❌ ReportLab not available. Install with: pip install reportlab")
                except Exception as e:
                    st.error(f"❌ Error creating bulk PDF: {str(e)}")
        
        with col3:
            if st.button("🗑️ Clear Current Analysis", key="clear_current_analysis", type="secondary"):
                st.session_state.current_ai_analysis = ""
                st.session_state.current_visualization = None
                st.rerun()
    
    # Analysis history
    if hasattr(st.session_state, 'insights_history') and st.session_state.insights_history:
        st.markdown("---")
        st.subheader("📚 Recent Analysis")
        
        for i, entry in enumerate(st.session_state.insights_history[-3:]):  # Show last 3
            with st.expander(f"🔍 {entry['question'][:50]}... ({entry['timestamp']})"):
                st.markdown(f"**Question:** {entry['question']}")
                st.markdown(f"**SQL Query:**")
                st.code(entry['sql_query'], language="sql")
                st.markdown(f"**AI Analysis:**")
                escaped_analysis = escape_markdown_for_streamlit(entry['analysis'])
                st.markdown(escaped_analysis)
                st.markdown(f"**Data Shape:** {entry['data_shape']}")
                
                # Export button for this specific insight
                if st.button(f"📄 Export as PDF", key=f"export_insight_{i}"):
                    try:
                        if REPORTLAB_AVAILABLE:
                            pdf_data = generate_pdf_report(entry, f"ai_insight_{entry['timestamp'].replace(':', '-').replace(' ', '_')}.pdf")
                            st.download_button(
                                label="📥 Download PDF",
                                data=pdf_data,
                                file_name=f"ai_insight_{entry['timestamp'].replace(':', '-').replace(' ', '_')}.pdf",
                                mime="application/pdf"
                            )
                        else:
                            st.error("❌ ReportLab not available. Install with: pip install reportlab")
                    except Exception as e:
                        st.error(f"❌ Error creating PDF: {str(e)}")
    


def custom_queries_tab():
    """Enhanced Custom Queries tab with AI generation and visualization"""
    st.header("💻 SQL Playground")
    
    if not st.session_state.loaded_tables and not st.session_state.data_processor.has_data():
        st.markdown('<div class="status-warning">⚠️ No data loaded. Please load data first.</div>', unsafe_allow_html=True)
        return
    
    # Get all available tables
    all_tables = {}
    if st.session_state.loaded_tables:
        all_tables.update(st.session_state.loaded_tables)
    if st.session_state.data_processor.has_data():
        all_tables["main"] = st.session_state.data_processor.get_data()
    
    if not all_tables:
        st.markdown('<div class="status-warning">⚠️ No tables available for SQL execution.</div>', unsafe_allow_html=True)
        return
    
    # Initialize query history if not exists
    if 'query_history' not in st.session_state:
        st.session_state.query_history = []
    
    # Display table overview
    st.subheader("📊 Available Tables")
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Tables", len(all_tables))
    
    with col2:
        total_rows = sum(df.shape[0] for df in all_tables.values())
        st.metric("Total Rows", f"{total_rows:,}")
    
    with col3:
        total_cols = sum(df.shape[1] for df in all_tables.values())
        st.metric("Total Columns", total_cols)
    
    with col4:
        executed_queries = len(st.session_state.query_history)
        st.metric("Executed Queries", executed_queries)
    
    # Table information
    with st.expander("📋 Table Details", expanded=False):
        for table_name, df in all_tables.items():
            st.markdown(f"**{table_name}** ({df.shape[0]} rows × {df.shape[1]} columns)")
            st.markdown(f"Columns: {', '.join(df.columns.tolist())}")
            st.markdown("---")
    
    # Main query interface
    st.subheader("✏️ SQL Query Editor")
    
    # SQL Query Editor
    current_query = st.session_state.get('custom_sql_query', '')
    if not current_query:
        # Pre-populate with a sample query
        sample_query = f"""-- Example queries for your data:
-- SELECT * FROM {list(all_tables.keys())[0]} LIMIT 10;
-- SELECT COUNT(*) as total_rows FROM {list(all_tables.keys())[0]};
-- SELECT * FROM {list(all_tables.keys())[0]} WHERE {list(all_tables[list(all_tables.keys())[0]].columns)[0]} IS NOT NULL;
"""
        current_query = sample_query
    
    custom_sql = st.text_area(
        "Write your SQL query:",
        value=current_query,
        height=200,
        placeholder="SELECT * FROM table_name LIMIT 10;",
        help="Write your own SQL query using the available tables",
        key="custom_sql_editor"
    )
    
    # Update current query in session state
    st.session_state.custom_sql_query = custom_sql
    
    # Quick query templates
    st.markdown("**📝 Quick Templates**")
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("🔍 Explore Data", key="explore_template_btn"):
            template = f"SELECT * FROM {list(all_tables.keys())[0]} LIMIT 10;"
            st.session_state.custom_sql_query = template
            st.rerun()
    
    with col2:
        if st.button("📊 Count Records", key="count_template_btn"):
            template = f"SELECT COUNT(*) as total_records FROM {list(all_tables.keys())[0]};"
            st.session_state.custom_sql_query = template
            st.rerun()
    
    with col3:
        if st.button("📈 Basic Stats", key="stats_template_btn"):
            first_table = list(all_tables.keys())[0]
            first_df = all_tables[first_table]
            
            # Find a numeric column for stats, or use COUNT only
            numeric_cols = first_df.select_dtypes(include=['number']).columns.tolist()
            
            if numeric_cols:
                # Use first numeric column for stats
                stats_col = numeric_cols[0]
                template = f"SELECT COUNT(*) as count, AVG({stats_col}) as avg_value, MIN({stats_col}) as min_value, MAX({stats_col}) as max_value FROM {first_table};"
            else:
                # No numeric columns, just use COUNT
                template = f"SELECT COUNT(*) as count, COUNT(DISTINCT {first_df.columns[0]}) as unique_values FROM {first_table};"
            
            st.session_state.custom_sql_query = template
            st.rerun()
    
    with col4:
        if st.button("🔗 Join Tables", key="join_template_btn") and len(all_tables) > 1:
            tables = list(all_tables.keys())
            table1, table2 = tables[0], tables[1]
            df1, df2 = all_tables[table1], all_tables[table2]
            
            # Find common columns for JOIN
            common_cols = set(df1.columns) & set(df2.columns)
            
            if common_cols:
                # Use first common column for JOIN
                join_col = list(common_cols)[0]
                template = f"SELECT * FROM {table1} t1 JOIN {table2} t2 ON t1.{join_col} = t2.{join_col} LIMIT 10;"
            else:
                # No common columns, create a simple cross join
                template = f"SELECT * FROM {table1} t1 CROSS JOIN {table2} t2 LIMIT 10;"
            
            st.session_state.custom_sql_query = template
            st.rerun()
    
    # Query execution controls
    col_exec1, col_exec2, col_exec3 = st.columns([2, 1, 1])
    
    with col_exec1:
        if st.button("▶️ Execute Query", key="execute_query_custom_btn", type="primary", use_container_width=True):
            if custom_sql and custom_sql.strip():
                with st.spinner("Executing SQL query..."):
                    try:
                        # Register all tables in DuckDB
                        conn = duckdb.connect()
                        for table_name, df in all_tables.items():
                            conn.register(table_name, df)
                        
                        # Execute query
                        result_df = conn.execute(custom_sql).fetchdf()
                        
                        if not result_df.empty:
                            st.success("✅ Query executed successfully!")
                            
                            # Store results and query in history
                            query_entry = {
                                'query': custom_sql,
                                'results': result_df,
                                'timestamp': pd.Timestamp.now(),
                                'rows_returned': len(result_df)
                            }
                            st.session_state.query_history.append(query_entry)
                            st.session_state.sql_playground_query_results = result_df
                            st.rerun()
                            
                        else:
                            st.warning("⚠️ Query returned no results.")
                            
                    except Exception as sql_error:
                        st.error(f"❌ SQL execution error: {str(sql_error)}")
                        st.info("💡 Check your SQL syntax and table/column names.")
                    finally:
                        conn.close()
            else:
                st.warning("⚠️ Please enter a SQL query.")
    
    with col_exec2:
        if st.button("🧹 Clear Query", key="clear_query_btn"):
            st.session_state.custom_sql_query = ""
            st.session_state.sql_playground_query_results = None
            st.rerun()
    
    with col_exec3:
        if st.button("📋 Copy Query", key="copy_query_btn"):
            st.code(custom_sql, language="sql")
    
    # Results and Visualization Section
    if 'sql_playground_query_results' in st.session_state and st.session_state.sql_playground_query_results is not None:
        st.markdown("---")
        st.subheader("📊 Query Results")
        
        result_df = st.session_state.sql_playground_query_results
        
        # Results summary
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Rows Returned", len(result_df))
        with col2:
            st.metric("Columns", len(result_df.columns))
        with col3:
            st.metric("Memory Usage", f"{result_df.memory_usage(deep=True).sum() / 1024:.1f} KB")
        
        # Display the results table
        st.markdown("**📋 Data Table**")
        safe_dataframe_display(result_df, width='stretch')
        
        # Visualization section
        if len(result_df) > 0:
            st.markdown("**📈 Visualization**")
            
            col_viz1, col_viz2 = st.columns([1, 1])
            
            with col_viz1:
                if st.button("🎨 Auto-Generate Visualization", key="auto_viz_custom_btn", type="primary", use_container_width=True):
                    try:
                        # Auto-detect the best visualization type based on data
                        numeric_cols = result_df.select_dtypes(include=[np.number]).columns
                        categorical_cols = result_df.select_dtypes(include=['object', 'category']).columns
                        
                        if len(numeric_cols) > 0 and len(categorical_cols) > 0:
                            # Bar chart for categorical vs numeric
                            fig = st.session_state.visualizer.create_custom_chart(
                                result_df, 
                                'bar', 
                                categorical_cols[0], 
                                numeric_cols[0],
                                title=f"{numeric_cols[0]} by {categorical_cols[0]}"
                            )
                            st.session_state.custom_viz_fig = fig
                            st.session_state.custom_viz_type = "Bar chart (categorical vs numeric)"
                        elif len(numeric_cols) > 1:
                            # Scatter plot for multiple numeric columns
                            fig = st.session_state.visualizer.create_custom_chart(
                                result_df, 
                                'scatter', 
                                numeric_cols[0], 
                                numeric_cols[1],
                                title=f"{numeric_cols[1]} vs {numeric_cols[0]}"
                            )
                            st.session_state.custom_viz_fig = fig
                            st.session_state.custom_viz_type = "Scatter plot (numeric correlation)"
                        elif len(categorical_cols) > 0:
                            # Pie chart for categorical data
                            # For pie chart, we need to count the categories
                            cat_counts = result_df[categorical_cols[0]].value_counts()
                            fig = st.session_state.visualizer.create_custom_chart(
                                cat_counts.reset_index(), 
                                'pie', 
                                categorical_cols[0], 
                                'count',
                                title=f"Distribution of {categorical_cols[0]}"
                            )
                            st.session_state.custom_viz_fig = fig
                            st.session_state.custom_viz_type = "Pie chart (categorical distribution)"
                        elif len(numeric_cols) == 1:
                            # Histogram for single numeric column
                            fig = st.session_state.visualizer.create_custom_chart(
                                result_df, 
                                'histogram', 
                                numeric_cols[0], 
                                numeric_cols[0],
                                title=f"Distribution of {numeric_cols[0]}"
                            )
                            st.session_state.custom_viz_fig = fig
                            st.session_state.custom_viz_type = "Histogram (numeric distribution)"
                        else:
                            # Try auto visualization as fallback
                            fig = st.session_state.visualizer.create_auto_visualization(result_df)
                            if fig:
                                st.session_state.custom_viz_fig = fig
                                st.session_state.custom_viz_type = "Auto-generated visualization"
                            else:
                                st.warning("ℹ️ Data structure not suitable for automatic visualization")
                            
                    except Exception as e:
                        st.error(f"❌ Could not generate visualization: {str(e)}")
                        st.info("💡 Try a different query or check your data structure")
            
            with col_viz2:
                if st.button("🧹 Clear Visualization", key="clear_viz_custom_btn", type="secondary", use_container_width=True):
                    if 'custom_viz_fig' in st.session_state:
                        del st.session_state.custom_viz_fig
                    if 'custom_viz_type' in st.session_state:
                        del st.session_state.custom_viz_type
                    st.rerun()
            
            # Display the visualization if available
            if 'custom_viz_fig' in st.session_state and st.session_state.custom_viz_fig:
                st.markdown(f"**{st.session_state.custom_viz_type}**")
                st.plotly_chart(st.session_state.custom_viz_fig, use_container_width=True)
    
    # Query History
    if st.session_state.query_history:
        st.markdown("---")
        st.subheader("📚 Query History")
        
        # Show recent queries
        recent_queries = st.session_state.query_history[-5:]  # Show last 5 queries
        
        for i, query_entry in enumerate(reversed(recent_queries)):
            with st.expander(f"Query {len(st.session_state.query_history) - i} - {query_entry['timestamp'].strftime('%H:%M:%S')} ({query_entry['rows_returned']} rows)"):
                st.code(query_entry['query'], language="sql")
                
                col_hist1, col_hist2 = st.columns([1, 1])
                with col_hist1:
                    if st.button(f"🔄 Re-run Query {len(st.session_state.query_history) - i}", key=f"rerun_query_{i}"):
                        st.session_state.custom_sql_query = query_entry['query']
                        st.rerun()
                
                with col_hist2:
                    if st.button(f"📊 Show Results {len(st.session_state.query_history) - i}", key=f"show_results_{i}"):
                        st.session_state.sql_playground_query_results = query_entry['results']
                        st.rerun()
        
        if st.button("🗑️ Clear History", key="clear_history_btn"):
            st.session_state.query_history = []
            st.rerun()
            if 'custom_sql_query' in st.session_state:
                del st.session_state.custom_sql_query
            if 'custom_sql_results' in st.session_state:
                del st.session_state.custom_sql_results
            st.rerun()
    
    # Navigation to next tab
    

def insights_tab():
    """Insights tab with formatted document view and PDF download"""
    st.header("📄 AI Insights & Reports")
    
    # Initialize insights_history if it doesn't exist
    if 'insights_history' not in st.session_state:
        st.session_state.insights_history = []
    
    if not st.session_state.insights_history:
        st.markdown('<div class="status-warning">⚠️ No insights generated yet. Please use the AI Analysis tab to generate insights first.</div>', unsafe_allow_html=True)
        return
    
    # Display all insights
    st.subheader("📚 Generated Insights")
    
    for i, entry in enumerate(st.session_state.insights_history):
        with st.expander(f"📊 Analysis #{i+1}: {entry['question'][:60]}... ({entry['timestamp']})", expanded=(i == len(st.session_state.insights_history) - 1)):
            
            # Create a formatted document view
            st.markdown("---")
            st.markdown(f"### 📝 **Question**")
            st.markdown(f"*{entry['question']}*")
            
            st.markdown(f"### 🔍 **SQL Query**")
            st.code(entry['sql_query'], language="sql")
            
            st.markdown(f"### 📊 **Data Summary**")
            st.markdown(f"- **Rows:** {entry['data_shape'][0]}")
            st.markdown(f"- **Columns:** {entry['data_shape'][1]}")
            st.markdown(f"- **Generated:** {entry['timestamp']}")
            
            st.markdown(f"### 🧠 **AI Analysis**")
            escaped_analysis = escape_markdown_for_streamlit(entry['analysis'])
            st.markdown(escaped_analysis)
            
            # Export buttons for this specific insight
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button(f"📄 Export as PDF", key=f"pdf_{i}"):
                    try:
                        if REPORTLAB_AVAILABLE:
                            # Generate proper PDF using ReportLab
                            pdf_data = generate_pdf_report(entry, f"ai_insights_report_{entry['timestamp'].replace(':', '-').replace(' ', '_')}.pdf")
                            st.download_button(
                                label="Download PDF Report",
                                data=pdf_data,
                                file_name=f"ai_insights_report_{entry['timestamp'].replace(':', '-').replace(' ', '_')}.pdf",
                                mime="application/pdf"
                            )
                        else:
                            st.error("❌ ReportLab not available. Install with: pip install reportlab")
                    except Exception as e:
                        st.error(f"❌ Error creating PDF report: {str(e)}")
            
            with col2:
                if st.button(f"📊 Export as Excel", key=f"excel_{i}"):
                    try:
                        # Create Excel file
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            # Analysis sheet
                            analysis_df = pd.DataFrame({
                                'Question': [entry['question']],
                                'SQL_Query': [entry['sql_query']],
                                'AI_Analysis': [entry['analysis']],
                                'Generated_At': [entry['timestamp']],
                                'Data_Rows': [entry['data_shape'][0]],
                                'Data_Columns': [entry['data_shape'][1]]
                            })
                            analysis_df.to_excel(writer, sheet_name='Analysis', index=False)
                        
                        st.download_button(
                            label="Download Excel Report",
                            data=output.getvalue(),
                            file_name=f"ai_insights_{entry['timestamp'].replace(':', '-').replace(' ', '_')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    except Exception as e:
                        st.error(f"❌ Error creating Excel report: {str(e)}")
            
            with col3:
                if st.button(f"📝 Export as Word", key=f"word_{i}"):
                    try:
                        if DOCX_AVAILABLE:
                            # Generate Word document
                            word_data = generate_word_report(entry)
                            st.download_button(
                                label="Download Word Report",
                                data=word_data,
                                file_name=f"ai_insights_{entry['timestamp'].replace(':', '-').replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.error("❌ python-docx not available. Install with: pip install python-docx")
                    except Exception as e:
                        st.error(f"❌ Error creating Word report: {str(e)}")
    
    # Additional export options
    if len(st.session_state.insights_history) > 0:
        st.markdown("---")
        st.subheader("📤 Bulk Export Options")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📄 Export All as PDF", key="bulk_pdf"):
                try:
                    if REPORTLAB_AVAILABLE:
                        # Create a combined PDF with all insights
                        buffer = io.BytesIO()
                        doc = SimpleDocTemplate(buffer, pagesize=A4)
                        story = []
                        
                        for i, entry in enumerate(st.session_state.insights_history):
                            if i > 0:
                                story.append(PageBreak())
                            
                            # Add title for each insight
                            title_style = ParagraphStyle('Title', fontSize=18, spaceAfter=20)
                            story.append(Paragraph(f"Analysis #{i+1}: {entry['question'][:50]}...", title_style))
                            
                            # Add content
                            pdf_data = generate_pdf_report(entry, "")
                            # Note: This is a simplified version - in practice, you'd need to parse the PDF content
                            story.append(Paragraph(f"Question: {entry['question']}", getSampleStyleSheet()['Normal']))
                            story.append(Paragraph(f"SQL: {entry['sql_query']}", getSampleStyleSheet()['Code']))
                            story.append(Paragraph(f"Analysis: {entry['analysis']}", getSampleStyleSheet()['Normal']))
                        
                        doc.build(story)
                        buffer.seek(0)
                        
                        st.download_button(
                            label="Download All Insights as PDF",
                            data=buffer.getvalue(),
                            file_name=f"all_insights_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf"
                        )
                    else:
                        st.error("❌ ReportLab not available. Install with: pip install reportlab")
                except Exception as e:
                    st.error(f"❌ Error creating bulk PDF: {str(e)}")
        
        with col2:
            if st.button("🗑️ Clear All Insights", type="secondary", key="clear_insights_btn_1"):
                st.session_state.insights_history = []
                st.rerun()
    
    # Clear all insights button
    if len(st.session_state.insights_history) > 0:
        st.markdown("---")
        if st.button("🗑️ Clear All Insights", type="secondary", key="clear_insights_btn_2"):
            st.session_state.insights_history = []
            st.rerun()
    

def settings_tab():
    """Settings and configuration tab"""
    st.header("⚙️ Settings & Configuration")
    
    # Create tabs within settings
    settings_tab1, settings_tab2, settings_tab3, settings_tab4 = st.tabs([
        "🤖 AI Models", 
        "🔗 Data Connectors", 
        "📤 Export Data", 
        "ℹ️ System Info"
    ])
    
    with settings_tab1:
        st.subheader("🤖 AI Model Management")
        
        if OLLAMA_AVAILABLE:
            available_models = get_available_models()
            missing_models = get_missing_models()
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.write(f"**Available Models:** {len(available_models)}")
                if available_models:
                    for model in available_models:
                        st.write(f"• {model}")
                else:
                    st.warning("No models found")
                
                # Show missing models
                if missing_models:
                    st.markdown("---")
                    st.write(f"**Missing Recommended Models:** {len(missing_models)}")
                    for model in missing_models[:5]:  # Show first 5
                        st.write(f"• {model}")
                    if len(missing_models) > 5:
                        st.write(f"... and {len(missing_models) - 5} more")
            
            with col2:
                if st.button("🔄 Refresh Model List", key="refresh_models_btn"):
                    st.rerun()
                
                if st.button("📥 Download Commands", key="download_commands_btn"):
                    if missing_models:
                        download_commands = generate_download_commands(missing_models)
                        st.code("\n".join(download_commands), language="bash")
                    else:
                        st.info("All recommended models are already available!")
                
                if st.button("⚙️ Open Model Setup", key="open_model_setup_btn"):
                    st.info("Run `python setup_models.py` in your terminal")
                
                # Debug section
                with st.expander("🔍 Debug Information", expanded=False):
                    if st.button("🔍 Test Model Detection", key="test_model_detection_btn"):
                        test_result = test_model_detection()
                        if test_result["status"] == "success":
                            st.success(f"✅ Model detection working! Found {test_result['models_found']} models.")
                            st.write("**Detected models:**")
                            for model in test_result["models"]:
                                st.write(f"• {model}")
                        else:
                            st.error(f"❌ Model detection failed: {test_result['error']}")
                        st.write(f"**Ollama Available:** {test_result['ollama_available']}")
        else:
            st.error("Ollama configuration not available")
    
    with settings_tab2:
        st.subheader("🔗 Data Connector Configuration")
        
        # Google Analytics Connector
        with st.expander("📊 Google Analytics Connector"):
            st.markdown("Connect to Google Analytics to import web traffic data")
            
            col1, col2 = st.columns(2)
            
            with col1:
                ga_property_id = st.text_input("Property ID", placeholder="123456789", key="ga_property_id")
                ga_start_date = st.date_input("Start Date", key="ga_start_date")
                ga_end_date = st.date_input("End Date", key="ga_end_date")
            
            with col2:
                ga_metrics = st.multiselect(
                    "Metrics to Import",
                    ["sessions", "users", "pageviews", "bounceRate", "avgSessionDuration"],
                    default=["sessions", "users", "pageviews"],
                    key="ga_metrics"
                )
                ga_dimensions = st.multiselect(
                    "Dimensions to Import", 
                    ["date", "country", "deviceCategory", "trafficSource"],
                    default=["date", "country"],
                    key="ga_dimensions"
                )
            
            if st.button("🔗 Test GA Connection", key="test_ga_btn"):
                st.info("GA connector test - requires authentication setup")
                st.code("""
# To set up GA connector:
# 1. Create Google Cloud project
# 2. Enable Analytics Reporting API
# 3. Create service account credentials
# 4. Download JSON key file
# 5. Configure in connectors/ga_connector.py
                """)
        
        # Snowflake Connector
        with st.expander("❄️ Snowflake Connector"):
            st.markdown("Connect to Snowflake data warehouse")
            
            col1, col2 = st.columns(2)
            
            with col1:
                sf_account = st.text_input("Account", placeholder="your-account.snowflakecomputing.com", key="sf_account")
                sf_username = st.text_input("Username", key="sf_username")
                sf_password = st.text_input("Password", type="password", key="sf_password")
            
            with col2:
                sf_database = st.text_input("Database", placeholder="PROD_DB", key="sf_database")
                sf_schema = st.text_input("Schema", placeholder="PUBLIC", key="sf_schema")
                sf_warehouse = st.text_input("Warehouse", placeholder="COMPUTE_WH", key="sf_warehouse")
            
            if st.button("🔗 Test Snowflake Connection", key="test_sf_btn"):
                st.info("Snowflake connector test - requires valid credentials")
                st.code("""
# To set up Snowflake connector:
# 1. Get account URL from Snowflake console
# 2. Create user with appropriate permissions
# 3. Configure connection parameters
# 4. Test connection in connectors/snowflake_connector.py
                """)
        
        # Oncore Connector
        with st.expander("🏥 Oncore Connector"):
            st.markdown("Connect to Oncore clinical data system")
            
            col1, col2 = st.columns(2)
            
            with col1:
                oncore_base_url = st.text_input("Base URL", placeholder="https://your-oncore-instance.com", key="oncore_base_url")
                oncore_username = st.text_input("Username", key="oncore_username")
                oncore_password = st.text_input("Password", type="password", key="oncore_password")
            
            with col2:
                oncore_database = st.text_input("Database", placeholder="ONCORE_PROD", key="oncore_database")
                oncore_study_id = st.text_input("Study ID (optional)", key="oncore_study_id")
            
            if st.button("🔗 Test Oncore Connection", key="test_oncore_btn"):
                st.info("Oncore connector test - requires valid credentials")
                st.code("""
# To set up Oncore connector:
# 1. Get Oncore instance URL
# 2. Create API user account
# 3. Configure database connection
# 4. Test connection in connectors/oncore_connector.py
                """)
        
        # Excel/CSV Connector (already working)
        with st.expander("📁 Excel/CSV Connector"):
            st.markdown("✅ Excel and CSV file upload is already configured and working")
            st.success("No additional configuration needed - use the Data Load tab to upload files")
    
    with settings_tab3:
        st.subheader("📤 Export Your Data")
        
        if st.session_state.loaded_tables or st.session_state.data_processor.has_data():
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("💾 Export as CSV", key="export_csv_btn"):
                    try:
                        # Get current data
                        if st.session_state.loaded_tables:
                            current_table = st.session_state.current_table if hasattr(st.session_state, 'current_table') else list(st.session_state.loaded_tables.keys())[0]
                            current_data = st.session_state.loaded_tables[current_table]
                        else:
                            current_data = st.session_state.data_processor.get_data()
                        
                        # Convert mixed data types to strings to avoid Arrow serialization issues
                        export_data = current_data.copy()
                        for col in export_data.columns:
                            if export_data[col].dtype == 'object':
                                export_data[col] = export_data[col].astype(str)
                            elif 'datetime' in str(export_data[col].dtype):
                                export_data[col] = export_data[col].dt.strftime('%Y-%m-%d %H:%M:%S')
                        
                        csv_data = export_data.to_csv(index=False)
                        st.download_button(
                            label="Download CSV",
                            data=csv_data,
                            file_name=f"business_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                            mime="text/csv"
                        )
                    except Exception as e:
                        st.error(f"Error exporting CSV: {str(e)}")
            
            with col2:
                if st.button("📊 Export All Tables", key="export_excel_btn"):
                    try:
                        # Export all loaded tables
                        export_data = {}
                        if st.session_state.loaded_tables:
                            export_data.update(st.session_state.loaded_tables)
                        if st.session_state.data_processor.has_data():
                            export_data["main"] = st.session_state.data_processor.get_data()
                        
                        # Create Excel file with multiple sheets
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            for sheet_name, df in export_data.items():
                                df.to_excel(writer, sheet_name=sheet_name, index=False)
                        
                        st.download_button(
                            label="Download Excel",
                            data=output.getvalue(),
                            file_name=f"all_tables_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    except Exception as e:
                        st.error(f"Error exporting Excel: {str(e)}")
        else:
            st.info("No data loaded to export")
    
    with settings_tab4:
        st.subheader("ℹ️ System Information")
        
        info_col1, info_col2 = st.columns(2)
        
        with info_col1:
            st.metric("Python Version", "3.10+")
            st.metric("Streamlit Version", "1.28.1")
            st.metric("DuckDB Version", "0.9.2")
        
        with info_col2:
            st.metric("Pandas Version", "2.1.3")
            st.metric("Ollama Integration", "Direct API")
            st.metric("Plotly Version", "5.17.0")
        
        # Additional system info
        st.markdown("---")
        st.subheader("🔧 Technical Details")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Available Connectors:**")
            st.write("• ✅ Excel/CSV (Active)")
            st.write("• 🔧 Google Analytics (Configurable)")
            st.write("• 🔧 Snowflake (Configurable)")
            st.write("• 🔧 Oncore (Configurable)")
        
        with col2:
            st.write("**System Status:**")
            st.write("• ✅ Offline Operation")
            st.write("• ✅ AI Analysis Ready")
            st.write("• ✅ Data Processing Active")
            st.write("• ✅ Visualization Ready")
    

if __name__ == "__main__":
    main()